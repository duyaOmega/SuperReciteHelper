#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""题库解析模块（由原单文件版本拆分）。"""

# 变量缩写约定（便于快速读懂复杂解析逻辑）：
# q / qs: question / question list（题目对象/题目列表）
# ans: answer（答案文本或答案集合）
# seg: segment（从样式、下划线或文本中提取的片段）
# nid / ilvl: docx 自动编号的 numId / 缩进层级
# rec: record（历史作答记录）
# m / pat: regex match / pattern（正则匹配对象/表达式）

import tkinter as tk
from tkinter import messagebox, font as tkfont, filedialog, ttk
import re
import json
import os
import random
import math
import hashlib
from datetime import datetime
import ctypes
import tempfile
import zipfile
import xml.etree.ElementTree as ET

ANSWER_LABEL_RE = re.compile(r'^(?:正确答案|答案|参考答案|标准答案|【答案】|\[答案\]|答|参考解答)\s*[：:]?\s*(.*)\s*$')

QUESTION_START_RE = re.compile(r'^\s*(?:\d{1,4}(?:[、．\)]|[.](?!\d))|[一二三四五六七八九十百零]+[.、．\)])\s*')

OPTION_PREFIX_RE = re.compile(r'^\s*([A-HＡ-Ｈ])[.、．,，\)）:：]\s*')

OPTION_TOKEN_RE = re.compile(r'([A-HＡ-Ｈ])[.、．,，\)）:：]\s*')

OPTION_MARKER_SYMBOLS = ('\uf0fe', '☑', '✅', '√', '✔', '☒', '✘', '✗')

SECTION_HEADING_PATTERNS = (
    ('single', re.compile(r'^\s*单选题\s*$')),
    ('multi', re.compile(r'^\s*多选题\s*$')),
    ('judge', re.compile(r'^\s*判断题\s*$')),
    ('blank', re.compile(r'^\s*填空题\s*$')),
    ('short', re.compile(r'^\s*简答题\s*$')),
)

QA_PROMPT_HINTS = (
    '什么是', '哪些', '哪几', '哪一', '如何', '为什么', '为何', '是否',
    '概念', '特点', '内容', '内涵', '作用', '意义', '影响', '要求',
    '区别', '分类', '趋势', '原则', '方法', '地位', '历程', '阶段',
    '名称', '含义', '实质', '时间', '对象', '标准', '功能', '战略',
)

def _clean_option_text(text):
    """清理选项中的版式残留符号（项目符号/勾选符号等）。"""
    t = str(text or '')
    t = re.sub(r'[\uf0b7\u2022]+', ' ', t)
    t = re.sub(r'[\uf0fe☑✅√✔☒✘✗]+', ' ', t)
    t = re.sub(r'\s{2,}', ' ', t)
    return t.strip()

def _option_contains_answer_marker(text):
    """判断选项文本是否包含“被选中/勾选”样式标记。"""
    t = str(text or '')
    return any(sym in t for sym in OPTION_MARKER_SYMBOLS)

def _extract_choice_answer_from_option_format(options):
    """从选项格式差异中提取答案（如带勾选符号的选项）。"""
    marked = []
    for k, v in options.items():
        if _option_contains_answer_marker(v):
            marked.append(k)
    return marked

def _detect_section_heading(text):
    """识别题库分区标题（单选/多选/判断/填空/简答）。"""
    t = (text or '').strip()
    for sec, pat in SECTION_HEADING_PATTERNS:
        if pat.match(t):
            return sec
    return None

def _extract_answer_keys_from_text(text):
    """从文末答案区提取题号->答案映射，支持分区（单选/多选/判断）。"""
    lines = [l.strip() for l in str(text or '').split('\n') if l.strip()]
    result = {
        'single': {},
        'multi': {},
        'judge': {},
        'blank': {},
        'short': {},
        'generic': {}
    }

    section = None
    started = False
    answer_section_title_re = re.compile(r'^\s*(?:(?:单选|多选|判断|填空|简答)题\s*)?(?:参考答案|标准答案|答案)\s*[：:]?\s*$')

    def is_answer_section_title(line):
        t = str(line or '').strip()
        if not t:
            return False
        if answer_section_title_re.match(t):
            return True

        # 兼容“程序设计实习 参考答案”这类带前缀标题。
        if re.search(r'(参考答案|标准答案|答案)\s*[：:]?\s*$', t):
            # 排除正文说明句，避免“文末附参考答案。”误触发。
            if any(k in t for k in ('文末附', '见文末', '附参考答案')):
                return False
            if any(p in t for p in ('。', '；', ';', '！', '？', '?')):
                return False
            return True

        return False
    # pair_re: 匹配“题号 + 答案token”写法，例如“12.A”或“12：正确”。
    pair_re = re.compile(r'(\d{1,4})\s*[.、:：]\s*([A-HＡ-Ｈ]{1,8}|正确|错误|对|错)')
    number_only_re = re.compile(r'^\s*(\d{1,4})\s*$')
    answer_only_letter_re = re.compile(r'^\s*([A-HＡ-Ｈ]{1,8})\s*$')
    answer_only_judge_re = re.compile(r'^\s*(正确|错误|对|错)\s*$')
    # pending_*: 处理“题号与答案分行书写”时的暂存队列，按顺序配对。
    pending_numbers = []
    pending_answers = []

    def flush_pending_to_current_section():
        if not pending_numbers or not pending_answers:
            return
        # “题号与答案分行”时，按出现顺序一一配对。
        pair_count = min(len(pending_numbers), len(pending_answers))
        for i in range(pair_count):
            qno = pending_numbers[i]
            ans = _normalize_answer_text(pending_answers[i])
            if section in result:
                result[section][qno] = ans
            result['generic'].setdefault(qno, ans)
        del pending_numbers[:pair_count]
        del pending_answers[:pair_count]

    for line in lines:
        # 进入答案区
        if is_answer_section_title(line):
            started = True
            continue

        if not started:
            continue

        # 分区切换（如“单选题参考答案”）
        if '单选' in line and '答案' in line:
            flush_pending_to_current_section()
            section = 'single'
            continue
        if '多选' in line and '答案' in line:
            flush_pending_to_current_section()
            section = 'multi'
            continue
        if '判断' in line and '答案' in line:
            flush_pending_to_current_section()
            section = 'judge'
            continue
        if '填空' in line and '答案' in line:
            flush_pending_to_current_section()
            section = 'blank'
            continue
        if '简答' in line and '答案' in line:
            flush_pending_to_current_section()
            section = 'short'
            continue

        # 新题区标题出现时退出答案分区
        sec_heading = _detect_section_heading(line)
        if sec_heading and '答案' not in line:
            flush_pending_to_current_section()
            section = None
            continue

        pairs = list(pair_re.finditer(line))

        for m in pairs:
            qno = int(m.group(1))
            ans = _normalize_answer_text(m.group(2))
            if section in result:
                result[section][qno] = ans
            result['generic'].setdefault(qno, ans)

        if pairs:
            continue

        # 兼容“题号与答案分行”的文末答案格式：
        # 例如：1\n2\n3\n...\nD\nC\nA
        m_num = number_only_re.match(line)
        if m_num:
            pending_numbers.append(int(m_num.group(1)))
            flush_pending_to_current_section()
            continue

        # 可能是单个字母答案（如 D / AC / Ｂ）
        m_ans_letter = answer_only_letter_re.match(line)
        if m_ans_letter:
            pending_answers.append(_normalize_answer_text(m_ans_letter.group(1)))
            flush_pending_to_current_section()
            continue

        # 判断题答案仅在判断题分区下采集，避免把正文“对/错”噪声误当答案。
        if section == 'judge':
            m_ans_judge = answer_only_judge_re.match(line)
            if m_ans_judge:
                pending_answers.append(_normalize_answer_text(m_ans_judge.group(1)))
                flush_pending_to_current_section()
                continue

        # 一行多个 token（空格/顿号/逗号分隔）
        # 仅在“整行几乎全是答案字符+分隔符”时采集，避免把正文（如 GPT）误当答案。
        if section == 'judge':
            multi_line_ok = bool(re.fullmatch(r'[A-HＡ-Ｈ正确错误对错\s,，、;；/\\|]+', line))
            token_pat = r'[A-HＡ-Ｈ]{1,8}|正确|错误|对|错'
        else:
            multi_line_ok = bool(re.fullmatch(r'[A-HＡ-Ｈ\s,，、;；/\\|]+', line))
            token_pat = r'[A-HＡ-Ｈ]{1,8}'

        if multi_line_ok:
            multi_tokens = re.findall(token_pat, line)
            if multi_tokens:
                pending_answers.extend(_normalize_answer_text(tok) for tok in multi_tokens)
                flush_pending_to_current_section()

    flush_pending_to_current_section()

    return result

def _parse_docx_numbered_choice_questions(filepath):
    """解析使用 Word 自动编号的选择题（题号/选项字母不在纯文本中）。"""
    try:
        from docx import Document
    except Exception:
        return []

    try:
        doc = Document(filepath)
    except Exception:
        return []

    def get_num_info(para):
        try:
            ppr = para._p.pPr if para._p is not None else None
            num_pr = ppr.numPr if ppr is not None else None
            if num_pr is None:
                return None, None
            # num_id: Word 自动编号实例 ID；ilvl: 同编号体系中的层级。
            num_id = int(num_pr.numId.val) if num_pr.numId is not None else None
            ilvl = int(num_pr.ilvl.val) if num_pr.ilvl is not None else None
            return num_id, ilvl
        except Exception:
            return None, None

    items = []
    for para in doc.paragraphs:
        text = (para.text or '').strip()
        if not text:
            continue

        if ('参考答案' in text) or ('标准答案' in text):
            break

        num_id, ilvl = get_num_info(para)
        items.append({'text': text, 'num_id': num_id, 'ilvl': ilvl})

    if not items:
        return []

    def looks_like_question_stem(text):
        t = str(text or '').strip()
        if not t:
            return False
        if OPTION_PREFIX_RE.match(t):
            return False
        if re.match(r'^[（(]\s*\d+\s*[)）]', t):
            return False
        if re.match(r'^[A-HＡ-Ｈ][.、．,，\)）:：]', t):
            return False
        if t.startswith('#include'):
            return False
        return len(t) >= 4

    # 识别“题干所在的主编号列表”：按结构特征评分而非单纯出现次数。
    num_id_counts = {}
    num_id_positions = {}
    for it in items:
        nid = it.get('num_id')
        if nid is None:
            continue
        num_id_counts[nid] = num_id_counts.get(nid, 0) + 1
        num_id_positions.setdefault(nid, []).append(it)

    if not num_id_counts:
        return []

    def score_num_id(nid):
        # rows: 属于同一 num_id 的全部段落，后续用启发式判断它像“题干编号”还是“选项编号”。
        rows = num_id_positions.get(nid, [])
        count = len(rows)
        stem_like = 0
        option_like = 0
        enum_like = 0
        punct_like = 0

        for row in rows:
            txt = str(row.get('text', '') or '').strip()
            if looks_like_question_stem(txt):
                stem_like += 1
            if OPTION_PREFIX_RE.match(txt) or re.match(r'^[A-HＡ-Ｈ][.、．,，\)）:：]', txt):
                option_like += 1
            if re.match(r'^[（(]\s*\d+\s*[)）]', txt):
                enum_like += 1
            if ('：' in txt) or (':' in txt) or ('？' in txt) or ('?' in txt):
                punct_like += 1

        # 高 stem_like + punct_like 倾向于题干编号；option/枚举倾向于选项或材料编号。
        # 分数越高，越倾向把该 num_id 作为“题干起始编号”。
        return (
            stem_like * 2.0 +
            punct_like * 0.8 +
            min(count, 30) * 0.05 -
            option_like * 1.5 -
            enum_like * 0.8
        )

    # question_num_id: 最像题干编号的 num_id，后续据此切分题块。
    question_num_id = max(num_id_counts.keys(), key=score_num_id)
    if num_id_counts.get(question_num_id, 0) < 8:
        return []

    question_starts = [
        idx for idx, it in enumerate(items)
        if it.get('num_id') == question_num_id
        and '试题结束' not in it.get('text', '')
        and '答案' not in it.get('text', '')
    ]

    if len(question_starts) < 8:
        return []

    def build_fallback_options(rows):
        # 当选项字母不可见时，按顺序映射 A/B/C/D...
        if not rows:
            return {}

        # 优先：块末尾若存在同一编号的连续段，通常就是选项列表。
        # tail_candidates: 题块末尾、同编号连续段的候选选项文本。
        tail_candidates = []
        idx = len(rows) - 1
        while idx >= 0 and not str(rows[idx].get('text', '') or '').strip():
            idx -= 1

        if idx >= 0:
            tail_num = rows[idx].get('num_id')
            if tail_num is not None and tail_num != question_num_id:
                j = idx
                while j >= 0:
                    rj = rows[j]
                    if rj.get('num_id') != tail_num:
                        break
                    txt = str(rj.get('text', '') or '').strip()
                    if txt:
                        tail_candidates.append(txt)
                    j -= 1
                tail_candidates.reverse()

        if 2 <= len(tail_candidates) <= 8:
            letters = 'ABCDEFGH'
            return {letters[i]: _clean_option_text(t) for i, t in enumerate(tail_candidates) if i < len(letters)}

        candidates = []
        for row in rows:
            t = str(row.get('text', '') or '').strip()
            if not t:
                continue
            if '参考答案' in t or '试题结束' in t:
                continue
            candidates.append(t)

        if len(candidates) > 8:
            candidates = candidates[-8:]

        if not (2 <= len(candidates) <= 8):
            return {}

        letters = 'ABCDEFGH'
        options = {}
        for i, t in enumerate(candidates):
            if i >= len(letters):
                break
            options[letters[i]] = _clean_option_text(t)
        return options

    questions = []
    for pos, start_idx in enumerate(question_starts):
        end_idx = question_starts[pos + 1] if pos + 1 < len(question_starts) else len(items)
        block_items = items[start_idx:end_idx]
        if not block_items:
            continue

        question_line = block_items[0].get('text', '').strip()
        block_lines = [it.get('text', '').strip() for it in block_items if it.get('text', '').strip()]
        if not question_line or len(block_lines) < 1:
            continue

        parsed = parse_single_block('\n'.join(block_lines))
        if parsed and parsed.get('options'):
            parsed['source_no'] = len(questions) + 1
            questions.append(parsed)
            continue

        fallback_options = build_fallback_options(block_items[1:])
        if fallback_options:
            questions.append({
                'id': 0,
                'text': question_line,
                'options': fallback_options,
                'answer': [],
                'type': 'single',
                'source_no': len(questions) + 1,
            })

    if not questions:
        return []

    extracted_text = extract_text_by_filetype(filepath)
    answer_keys = _extract_answer_keys_from_text(extracted_text)
    generic_map = answer_keys.get('generic', {})

    for idx, q in enumerate(questions, 1):
        token = generic_map.get(idx, '')
        if not token:
            continue

        letters = _extract_choice_answer(token, q.get('options') or {})
        if not letters:
            normalized = _normalize_answer_text(token)
            letters = re.findall(r'[A-H]', normalized)
            letters = sorted(set(letters), key=letters.index) if letters else []

        q['answer'] = letters
        if _is_judge_options(q.get('options') or {}):
            q['type'] = 'judge'
        elif len(letters) > 1:
            q['type'] = 'multi'
        else:
            q['type'] = 'single'

    for idx, q in enumerate(questions, 1):
        q['id'] = idx

    return questions

def _fill_answers_from_answer_keys(questions, answer_keys):
    """把文末答案区映射回题目列表。"""
    section_index = {'single': 0, 'multi': 0, 'judge': 0}

    for q in questions:
        options = q.get('options') or {}
        if not options:
            continue

        # 已有答案则跳过
        if q.get('answer'):
            continue

        section_hint = q.get('section_hint')
        ans_token = ''

        if section_hint in ('single', 'multi', 'judge'):
            section_index[section_hint] += 1
            ans_token = answer_keys.get(section_hint, {}).get(section_index[section_hint], '')

        if not ans_token:
            src_no = q.get('source_no')
            if src_no is not None:
                ans_token = answer_keys.get('generic', {}).get(src_no, '')

        if not ans_token:
            continue

        letters = _extract_choice_answer(ans_token, options)
        if not letters:
            letters = re.findall(r'[A-H]', _normalize_answer_text(ans_token))

        if letters:
            q['answer'] = sorted(set(letters), key=letters.index)
            if section_hint == 'multi':
                q['type'] = 'multi'
            elif section_hint == 'judge':
                q['type'] = 'judge'
            elif section_hint == 'single' and q.get('type') != 'judge':
                q['type'] = 'single'

def _normalize_extracted_text(text):
    """规范化不同来源文本，提升题块边界与选项识别稳定性。"""
    if not text:
        return ''

    t = text.replace('\r\n', '\n').replace('\r', '\n')
    t = t.replace('\u3000', ' ')
    # 标准答案标签前强制换行
    t = re.sub(r'(?<!\n)(正确答案|参考答案|标准答案|答案|答)[：:]', r'\n\1：', t)
    # 题号前强制换行，兼容 OCR/PDF 黏连；避免把小数(如4.5)误判成题号。
    # 仅在“前一字符不是数字”时插入换行，避免把 10. / 11. 误拆成 1\n0. / 1\n1.
    t = re.sub(r'(?<![\n\d])(\d{1,4}(?:[、．\)]|[.](?!\d)))', r'\n\1', t)
    t = re.sub(r'\n{3,}', '\n\n', t)
    return t.strip()

def _looks_like_option_line(text):
    if not text:
        return False
    if OPTION_PREFIX_RE.match(text):
        return True
    markers = OPTION_TOKEN_RE.findall(text)
    return len(markers) >= 2

def _has_option_structure(lines):
    """跨多行判断是否包含客观题选项结构。"""
    if not lines:
        return False

    prefixed = sum(1 for line in lines if OPTION_PREFIX_RE.match(line))
    if prefixed >= 2:
        return True

    joined = ' '.join(lines)
    markers = OPTION_TOKEN_RE.findall(joined)
    return len(markers) >= 2

def _looks_like_question_start_line(text):
    """判断一行是否像“新题开始”，避免把代码行（如 `0) {}`）误判成题号。"""
    t = (text or '').strip()
    if not QUESTION_START_RE.match(t):
        return False

    # 题号后的主体内容用于进一步判别。
    body = QUESTION_START_RE.sub('', t, count=1).strip()
    if not body:
        return False

    has_cjk = bool(re.search(r'[\u4e00-\u9fff]', body))
    has_q_hint = any(k in body for k in ('下列', '以下', '哪个', '哪项', '何者', '正确', '错误', '？', '?', '（ ）', '()'))

    # 代码样式行：常见于 C/C++/Java 片段，不应触发新题分块。
    if any(sym in body for sym in ('{', '}', ';', '#include', '::')) and not (has_cjk and has_q_hint):
        return False
    if re.search(r'\b(?:int|void|class|return)\b', body) and not has_cjk:
        return False

    # 仅含极短数字/符号片段通常是代码残片（如“0) { } };”）。
    if not has_cjk:
        compact = re.sub(r'\s+', '', body)
        if len(compact) <= 12 and re.search(r'[{};=<>]', compact):
            return False

    return True

def _looks_like_answer_token(text):
    """判断文本是否像标准答案（而不是选项行）。"""
    if not text:
        return False
    if _looks_like_option_line(text):
        return False

    normalized = _normalize_answer_text(text)
    cleaned = re.sub(r'[\s,，、;；/\\]+', '', normalized)
    if re.fullmatch(r'[A-H]+', cleaned or ''):
        return True

    answer_keywords = ('正确', '错误', '对', '错', '是', '否')
    return any(k == text.strip() for k in answer_keywords)

def _read_text_file(filepath):
    for enc in ('utf-8-sig', 'utf-8', 'utf-16', 'utf-16-le', 'utf-16-be', 'gb18030', 'gbk'):
        try:
            with open(filepath, 'r', encoding=enc) as f:
                return f.read()
        except Exception:
            continue
    raise ValueError(f'无法读取文本文件编码：{filepath}')

def _extract_docx_text_with_style(filepath):
    """优先用 python-docx 提取，尽量利用字体样式识别答案。"""
    try:
        from docx import Document
    except Exception:
        return None

    def run_is_styled(run):
        font = run.font
        color = None
        if font and font.color is not None and font.color.rgb is not None:
            color = str(font.color.rgb)
        return bool(
            run.bold or run.italic or run.underline or
            (font and font.highlight_color is not None) or
            color
        )

    lines = []
    doc = Document(filepath)

    def consume_para(para):
        raw = ''.join(r.text for r in para.runs).strip()
        if not raw:
            return
        styled = ''.join(r.text for r in para.runs if run_is_styled(r)).strip()
        lines.append(raw)

        # 若答案仅以不同字体标记，则注入标准答案行
        if styled and styled != raw and (QUESTION_START_RE.match(raw) or _is_blank_question(raw)):
            lines.append(f'答案：{styled}')

    for para in doc.paragraphs:
        consume_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    consume_para(para)

    return _normalize_extracted_text('\n'.join(lines).strip())

def _docx_run_is_red(run):
    """判断 docx run 是否为红色（兼容 RGB 与底层 XML 颜色值）。"""
    try:
        color = run.font.color if run.font else None
        if color and color.rgb:
            rgb = str(color.rgb).upper()
            if rgb.startswith('FF') and rgb[2:4] in ('00', '11', '22', '33') and rgb[4:6] in ('00', '11', '22', '33'):
                return True

        color_el = run._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
        if color_el is not None:
            val = (color_el.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') or '').upper().strip('#')
            if val in {'FF0000', 'C00000', 'CC0000', 'DD0000', 'EE0000'}:
                return True
    except Exception:
        return False
    return False

def _docx_run_is_nonblack(run):
    """判断 run 是否为非黑色字体（含主题色/显式颜色）。"""

    def is_near_black(rgb_hex):
        if not rgb_hex or len(rgb_hex) != 6:
            return False
        try:
            r = int(rgb_hex[0:2], 16)
            g = int(rgb_hex[2:4], 16)
            b = int(rgb_hex[4:6], 16)
        except Exception:
            return False

        # 深灰正文（如 333333）不应被视作“强调色”。
        return max(r, g, b) <= 70 and (max(r, g, b) - min(r, g, b)) <= 12

    try:
        color = run.font.color if run.font else None
        if color is None:
            return False

        if color.rgb is not None:
            rgb = str(color.rgb).upper()
            rgb = rgb.replace('#', '')
            if len(rgb) == 8:
                rgb = rgb[-6:]
            if rgb and rgb not in {'000000', '00000000'} and not is_near_black(rgb):
                return True
    except Exception:
        return False
    return False

def _docx_has_nondefault_shading(element):
    """检测 XML 节点上的底纹/背景色是否为非默认值。"""
    if element is None:
        return False

    try:
        shd_nodes = element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
    except Exception:
        shd_nodes = []

    for shd in shd_nodes:
        fill = (shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill') or '').upper()
        color = (shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color') or '').upper()
        val = (shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') or '').lower()

        # 空值、auto、白色视作默认背景；其余颜色或样式视作有背景标记。
        if fill and fill not in {'AUTO', 'FFFFFF', '00000000'}:
            return True
        if color and color not in {'AUTO', 'FFFFFF', '000000', '00000000'}:
            return True
        if val and val not in {'clear', 'nil'}:
            return True

    return False

def _docx_run_is_emphasis(run):
    """用于填空提取：下划线/加粗/非黑色/高亮均视为强调。"""
    if _docx_run_is_nonblack(run):
        return True
    if bool(getattr(run, 'underline', False)):
        return True
    if bool(getattr(run, 'bold', False)):
        return True
    try:
        if run.font and run.font.highlight_color is not None:
            return True
    except Exception:
        pass

    # 支持底层 XML 的背景色/底纹（某些文档把答案标在背景色而不是高亮里）。
    try:
        if _docx_has_nondefault_shading(run._element):
            return True
        parent = getattr(run._element, 'getparent', lambda: None)()
        if parent is not None and _docx_has_nondefault_shading(parent):
            return True
    except Exception:
        pass

    return False

def _parse_docx_styled_blank_questions(filepath):
    """从 docx 样式中提取填空题：强调样式文本视为答案片段。"""
    try:
        from docx import Document
    except Exception:
        return []

    try:
        doc = Document(filepath)
    except Exception:
        return []

    def iter_paragraphs():
        for p in doc.paragraphs:
            yield p
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    questions = []
    seen_texts = set()

    for para in iter_paragraphs():
        raw = para.text.strip()
        if not raw:
            continue
        if _looks_like_option_line(raw):
            continue

        full_text = ''.join((r.text or '') for r in para.runs).strip()
        if not full_text:
            full_text = raw

        if len(full_text) < 12 or len(full_text) > 260:
            continue
        if not any(ch in full_text for ch in ('，', '。', '；', ';', '：', ':', ',')):
            continue
        if _has_option_structure([full_text]):
            continue

        # 提取连续强调片段
        segments = []
        current = []
        for run in para.runs:
            t = (run.text or '')
            if not t.strip():
                continue
            if _docx_run_is_emphasis(run):
                current.append(t)
            else:
                if current:
                    seg = ''.join(current).strip(' ，,。；;：:、()（）[]【】')
                    if seg:
                        segments.append(seg)
                    current = []
        if current:
            seg = ''.join(current).strip(' ，,。；;：:、()（）[]【】')
            if seg:
                segments.append(seg)

        # 去重并过滤噪声
        cleaned_segments = []
        seen_seg = set()
        for seg in segments:
            seg = re.sub(r'\s+', '', seg)
            if not seg or seg in seen_seg:
                continue
            if len(seg) > 40:
                continue
            if re.fullmatch(r'[\W_]+', seg):
                continue
            if not re.search(r'[A-Za-z0-9\u4e00-\u9fff]', seg):
                continue
            seen_seg.add(seg)
            cleaned_segments.append(seg)

        if not cleaned_segments:
            continue

        # 过高强调占比多为整句加粗标题，不当作填空。
        ratio = sum(len(s) for s in cleaned_segments) / max(len(full_text), 1)
        if ratio < 0.03 or ratio > 0.55:
            continue
        if len(cleaned_segments) > 10:
            continue

        question_text = full_text
        answers = []
        replaced = False
        for idx, seg in enumerate(cleaned_segments, 1):
            placeholder = f'（{idx}）______'
            if seg in question_text:
                question_text = question_text.replace(seg, placeholder, 1)
                answers.append(f'（{idx}）{seg}')
                replaced = True

        if not replaced:
            continue

        question_text = QUESTION_START_RE.sub('', question_text).strip()
        if question_text in seen_texts:
            continue
        seen_texts.add(question_text)

        questions.append({
            'id': 0,
            'text': question_text,
            'options': {},
            'answer': _clean_answer_text('；'.join(answers)),
            'type': 'blank'
        })

    for idx, q in enumerate(questions, 1):
        q['id'] = idx
    return questions

def _merge_docx_blank_questions(base_questions, blank_questions):
    """将样式填空题并入解析结果，避免重复题干。"""
    if not blank_questions:
        return base_questions

    def normalize_for_match(text):
        t = (text or '')
        t = re.sub(r'（\s*\d+\s*）\s*[_＿﹍]+', ' ', t)
        t = re.sub(r'[_＿﹍]+', ' ', t)
        t = re.sub(r'[\s，,。；;：:、（）()\[\]【】]+', '', t)
        return t

    blank_anchors = []
    for b in blank_questions:
        txt = b.get('text', '')
        parts = re.split(r'（\s*\d+\s*）\s*[_＿﹍]+', txt)
        anchors = [normalize_for_match(p) for p in parts if normalize_for_match(p)]
        blank_anchors.append(anchors)

    filtered_base = []
    for q in base_questions:
        if q.get('options'):
            filtered_base.append(q)
            continue

        q_text_norm = normalize_for_match(q.get('text', ''))
        should_drop = False
        if q.get('type') in ('short', 'blank') and q_text_norm:
            for anchors in blank_anchors:
                if not anchors:
                    continue
                hit_count = sum(1 for a in anchors if len(a) >= 6 and a in q_text_norm)
                long_hit = any(len(a) >= 12 and a in q_text_norm for a in anchors)
                if long_hit or hit_count >= 2:
                    should_drop = True
                    break

        if not should_drop:
            filtered_base.append(q)

    merged = list(filtered_base)
    existing = set()
    for q in merged:
        t = re.sub(r'\s+', '', q.get('text', ''))
        if t:
            existing.add(t)

    for b in blank_questions:
        t = re.sub(r'\s+', '', b.get('text', ''))
        if not t or t in existing:
            continue
        merged.append(dict(b))
        existing.add(t)

    for idx, q in enumerate(merged, 1):
        q['id'] = idx
    return merged

def _merge_docx_preferred_with_short_questions(preferred_questions, parsed_questions):
    """保留 DOCX 高置信客观题，同时补回通用解析提取到的简答题。"""
    if not preferred_questions:
        return parsed_questions

    merged = [dict(q) for q in preferred_questions]

    def looks_like_valid_short_question(q):
        q_text = str(q.get('text', '') or '').strip()
        q_answer = str(q.get('answer', '') or '').strip()
        source_no = q.get('source_no')

        # 题干过短或仅标点通常是误拆分噪声（如 "?"、"::"）。
        if len(re.sub(r'\s+', '', q_text)) < 6:
            return False
        if not re.search(r'[A-Za-z\u4e00-\u9fff0-9]', q_text):
            return False

        # 兜底要求：要么带题号来源，要么题干明确是提问语句。
        if source_no is None and not (('？' in q_text) or ('?' in q_text) or any(h in q_text for h in QA_PROMPT_HINTS)):
            return False

        if not q_answer:
            return False

        ans_lines = [l.strip() for l in q_answer.split('\n') if l.strip()]
        stem_like_lines = sum(1 for l in ans_lines if ('下列' in l or '以下' in l) and (l.endswith('：') or l.endswith(':') or l.endswith('？') or l.endswith('?')))
        option_like_lines = sum(1 for l in ans_lines if _looks_like_option_line(l))
        qstart_like_lines = sum(1 for l in ans_lines if _looks_like_question_start_line(l))

        # 若“答案”里出现大量题干/题号/选项结构，通常是整段客观题被误吞，不应并入简答。
        polluted = (
            (len(ans_lines) >= 8 and (stem_like_lines >= 3 or option_like_lines >= 4 or qstart_like_lines >= 2)) or
            (len(re.sub(r'\s+', '', q_answer)) >= 320 and (stem_like_lines >= 2 or qstart_like_lines >= 2))
        )
        if polluted:
            return False

        return True

    def normalize_short_key(text):
        t = re.sub(r'\s+', '', str(text or ''))
        t = re.sub(r'[。．.!！?？]+$', '', t)
        return t

    seen_keys = set()
    for q in merged:
        key = normalize_short_key(q.get('text', ''))
        if key:
            seen_keys.add(key)

    for q in parsed_questions:
        if q.get('options'):
            continue
        if q.get('type') != 'short':
            continue
        if not looks_like_valid_short_question(q):
            continue

        key = normalize_short_key(q.get('text', ''))
        if not key or key in seen_keys:
            continue

        merged.append(dict(q))
        seen_keys.add(key)

    for idx, q in enumerate(merged, 1):
        q['id'] = idx
    return merged

def _mask_blank_question_text(question_text, answer_text):
    """兜底：若填空题题干未挖空，则根据答案片段自动替换为空位。"""
    text = (question_text or '').strip()
    if not text:
        return text
    if _is_blank_question(text):
        return text

    ans = (answer_text or '').strip()
    if not ans:
        return text

    # 优先匹配“（1）答案；（2）答案”格式。
    pairs = re.findall(r'（\s*(\d+)\s*）\s*([^；;\n]+)', ans)
    if pairs:
        out = text
        replaced = False
        for idx, (_, seg) in enumerate(pairs, 1):
            seg = seg.strip()
            if seg and seg in out:
                out = out.replace(seg, f'（{idx}）______', 1)
                replaced = True
        if replaced:
            return out

    # 次级策略：分号切分答案片段后替换。
    chunks = [c.strip() for c in re.split(r'[；;]+', ans) if c.strip()]
    out = text
    replaced = 0
    for i, chunk in enumerate(chunks, 1):
        chunk = re.sub(r'^（\s*\d+\s*）', '', chunk).strip()
        if chunk and chunk in out:
            out = out.replace(chunk, f'（{i}）______', 1)
            replaced += 1

    return out if replaced > 0 else text

def build_parse_candidates(filepath):
    """构建题库解析候选结果，供预览时切换比对。"""
    candidates = []
    ext = os.path.splitext(filepath)[1].lower()

    auto_questions = parse_questions(filepath)
    candidates.append(('自动识别（推荐）', auto_questions, '综合策略自动选择并合并结果'))

    if ext == '.docx':
        red_questions = _parse_docx_questions_with_red(filepath)
        if red_questions:
            candidates.append(('仅红色选项识别', red_questions, '按红色字体识别客观题答案'))

        styled_blanks = _parse_docx_styled_blank_questions(filepath)
        if styled_blanks:
            candidates.append(('仅样式填空识别', styled_blanks, '将下划线/非黑色/加粗视为填空答案'))

        if red_questions and styled_blanks:
            merged = _merge_docx_blank_questions(red_questions, styled_blanks)
            candidates.append(('红色客观 + 样式填空', merged, '客观题用红色，填空题用样式提取'))

    if ext == '.pdf':
        forced_judge = _coerce_no_option_questions_to_judge(auto_questions)
        if forced_judge:
            candidates.append(('强制判断题模式', forced_judge, '将无选项题按判断题处理（A=正确，B=错误）'))

    unique = []
    seen = set()
    for name, qs, desc in candidates:
        signature = (len(qs), tuple(sorted((q.get('type', ''), len(q.get('options', {}))) for q in qs[:50])))
        key = (name, signature)
        if key in seen:
            continue
        seen.add(key)
        unique.append((name, qs, desc))
    return unique

def _docx_extract_option_segments(para_text):
    """从段落文本中提取选项段：(letter, start, end, option_text)。"""
    segments = []
    marker_re = re.compile(r'([A-HＡ-Ｈ])[.、．,，\)）:：]\s*')
    markers = list(marker_re.finditer(para_text))
    if not markers:
        return segments

    for i, m in enumerate(markers):
        letter = m.group(1).translate(str.maketrans('ＡＢＣＤＥＦＧＨ', 'ABCDEFGH'))
        start = m.start()
        content_start = m.end()
        content_end = markers[i + 1].start() if i + 1 < len(markers) else len(para_text)
        option_text = para_text[content_start:content_end].strip().rstrip('。.，,；;')
        segments.append((letter, start, content_end, option_text))
    return segments

def _parse_docx_questions_with_red(filepath):
    """直接解析 docx：按红色选项识别正确答案，避免文本回退误判。"""
    try:
        from docx import Document
    except Exception:
        return []

    try:
        doc = Document(filepath)
    except Exception:
        return []

    def iter_paragraphs():
        for p in doc.paragraphs:
            yield p
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    questions = []
    current = None

    def finalize_current():
        nonlocal current
        if not current:
            return

        question_text = current['text']
        question_text = re.sub(r'^[\d]+[.、．\s]+', '', question_text)
        question_text = re.sub(r'【[^】]*】\s*', '', question_text)
        question_text = question_text.strip()

        if not question_text or not current['options']:
            current = None
            return

        red_answers = sorted(current['red_options'])
        if not red_answers and current['explicit_answer']:
            red_answers = _extract_choice_answer(current['explicit_answer'], current['options'])

        if _is_judge_options(current['options']):
            q_type = 'judge'
        elif len(red_answers) > 1:
            q_type = 'multi'
        else:
            q_type = 'single'

        questions.append({
            'id': 0,
            'text': question_text,
            'options': current['options'],
            'answer': red_answers,
            'type': q_type
        })
        current = None

    for para in iter_paragraphs():
        para_text = para.text.strip()
        if not para_text:
            continue

        if QUESTION_START_RE.match(para_text):
            finalize_current()
            current = {
                'text': para_text,
                'options': {},
                'red_options': set(),
                'explicit_answer': ''
            }
            continue

        if current is None:
            continue

        ans_match = ANSWER_LABEL_RE.match(para_text)
        if ans_match:
            current['explicit_answer'] = ans_match.group(1).strip()
            continue

        full_text = ''.join((r.text or '') for r in para.runs)
        if not full_text.strip():
            full_text = para_text

        segments = _docx_extract_option_segments(full_text)
        if not segments:
            # 续行题干
            if not current['options']:
                current['text'] = (current['text'] + ' ' + para_text).strip()
            continue

        # 构建字符级红色掩码，按“选项段”判断哪一项为红色。
        red_mask = []
        for run in para.runs:
            txt = run.text or ''
            if not txt:
                continue
            red_mask.extend([_docx_run_is_red(run)] * len(txt))
        if len(red_mask) < len(full_text):
            red_mask.extend([False] * (len(full_text) - len(red_mask)))

        for letter, start, end, option_text in segments:
            if option_text:
                current['options'][letter] = option_text

            # 仅检测本选项的有效区间：去掉尾部空白，避免“下个选项前红色空格”误伤当前选项。
            detect_end = end
            while detect_end > start and full_text[detect_end - 1].isspace():
                detect_end -= 1

            # 该选项有效区间任一字符为红色，即视为正确选项。
            if detect_end > start and any(red_mask[start:detect_end]):
                current['red_options'].add(letter)

    finalize_current()

    for idx, q in enumerate(questions, 1):
        q['id'] = idx
    return questions

def _extract_docx_text_fallback(filepath):
    """不依赖第三方库的 docx 文本提取回退方案。"""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    with zipfile.ZipFile(filepath, 'r') as zf:
        xml_data = zf.read('word/document.xml')

    root = ET.fromstring(xml_data)
    lines = []
    for p in root.findall('.//w:p', ns):
        texts = []
        for t in p.findall('.//w:t', ns):
            texts.append(t.text or '')
        line = ''.join(texts).strip()
        if line:
            lines.append(line)
    return _normalize_extracted_text('\n'.join(lines).strip())

def _extract_doc_text_windows(filepath):
    """使用 Windows Word COM 将 .doc/.docx 转成 txt（若本机可用）。"""
    try:
        import win32com.client  # type: ignore
    except Exception:
        return None

    temp_txt = None
    word = None
    doc = None
    try:
        fd, temp_txt = tempfile.mkstemp(suffix='.txt')
        os.close(fd)

        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(filepath))
        # 7 = wdFormatUnicodeText
        doc.SaveAs(os.path.abspath(temp_txt), FileFormat=7)
        doc.Close(False)
        doc = None
        word.Quit()
        word = None
        return _normalize_extracted_text(_read_text_file(temp_txt))
    except Exception:
        return None
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        if temp_txt and os.path.exists(temp_txt):
            try:
                os.remove(temp_txt)
            except Exception:
                pass

def _pdf_span_is_styled(span):
    font_name = str(span.get('font', '')).lower()
    color = span.get('color', 0)
    flags = int(span.get('flags', 0))
    if any(k in font_name for k in ('bold', 'black', 'heavy', 'demi', 'medium')):
        return True
    if color not in (0, None):
        return True
    # 低位标志中某些位常用于强调样式，作为启发式
    if flags & 2 or flags & 16:
        return True
    return False

def _extract_styled_segments_from_spans(spans):
    """从PDF行内span提取连续样式片段（红字/下划线/粗体等）。"""
    def style_signature(span):
        # style_signature: 用字体名/字号/flags/颜色构成样式指纹，便于比较“是否同样式”。
        return (
            str(span.get('font', '')),
            round(float(span.get('size', 0) or 0), 2),
            int(span.get('flags', 0) or 0),
            int(span.get('color', 0) or 0)
        )

    valid_spans = [s for s in spans if (s.get('text', '') or '').strip()]
    if not valid_spans:
        return []

    # 计算本行“主样式”：按文本长度加权出现最多的样式。
    # sig_weight: 不同样式的“文本长度加权”统计，用于估计一行的主样式。
    sig_weight = {}
    for s in valid_spans:
        sig = style_signature(s)
        sig_weight[sig] = sig_weight.get(sig, 0) + len((s.get('text', '') or '').strip())
    # dominant_sig: 该行最常见样式（通常是正文样式）。
    dominant_sig = max(sig_weight.items(), key=lambda kv: kv[1])[0]

    def is_contrasted_styled(span):
        # 仅将“有强调样式且不同于主样式”的片段视为答案候选，降低整行误判。
        return _pdf_span_is_styled(span) and style_signature(span) != dominant_sig

    segments = []
    current = []

    for s in valid_spans:
        txt = (s.get('text', '') or '').strip()
        if not txt:
            continue

        if is_contrasted_styled(s):
            current.append(txt)
        else:
            if current:
                seg = ''.join(current).strip(' ，,。；;：:、')
                if seg:
                    segments.append(seg)
                current = []

    if current:
        seg = ''.join(current).strip(' ，,。；;：:、')
        if seg:
            segments.append(seg)

    # 若完全没有“差异样式”，则回退为空，避免把整行普通文本误当答案。
    if not segments:
        return []

    # 去重（保持顺序）
    seen = set()
    unique = []
    for seg in segments:
        if seg in seen:
            continue
        # 过滤明显噪声
        if len(seg) < 2 and not re.fullmatch(r'[\u4e00-\u9fff]', seg):
            continue
        if re.fullmatch(r'[\W_]+', seg):
            continue
        seen.add(seg)
        unique.append(seg)
    return unique

def _dedupe_keep_order(items):
    seen = set()
    out = []
    for x in items:
        if x in seen:
            continue
        seen.add(x)
        out.append(x)
    return out

def _normalize_pdf_sentence(text):
    t = text or ''
    # 修复中文文本因换行导致的断词空格
    t = re.sub(r'(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])', '', t)
    t = re.sub(r'\s{2,}', ' ', t)
    return t.strip()

def _merge_split_segments(paragraph_text, segments):
    """合并被换行拆开的相邻答案片段（如“习近平新”+“时代...”）。"""
    if not segments:
        return segments

    merged = []
    i = 0
    while i < len(segments):
        current = segments[i]
        j = i + 1
        while j < len(segments):
            candidate = current + segments[j]
            if candidate in paragraph_text:
                current = candidate
                j += 1
            else:
                break
        merged.append(current)
        i = j

    return _dedupe_keep_order(merged)

def _extract_pdf_underline_rects(page):
    """提取 PDF 页面中的细横线矩形（常见于下划线标注答案）。"""
    rects = []
    try:
        drawings = page.get_drawings()
    except Exception:
        return rects

    for d in drawings:
        for it in d.get('items', []):
            if not it or it[0] != 're':
                continue
            r = it[1]
            x0, y0, x1, y1 = float(r.x0), float(r.y0), float(r.x1), float(r.y1)
            w = max(0.0, x1 - x0)
            h = max(0.0, y1 - y0)

            # 下划线通常是“很薄但较长”的横向矩形。
            if w >= 12 and h <= 1.2:
                rects.append((x0, y0, x1, y1))
    return rects

def _build_pdf_raw_line_records(raw_data):
    """从 rawdict 构建行级字符记录，供下划线字符级匹配。"""
    records = []
    for block in raw_data.get('blocks', []):
        if block.get('type') != 0:
            continue
        for line in block.get('lines', []):
            bbox = line.get('bbox')
            chars = []
            for span in line.get('spans', []):
                for ch in span.get('chars', []):
                    c = ch.get('c', '')
                    cb = ch.get('bbox')
                    if not c or not cb or len(cb) < 4:
                        continue
                    chars.append((c, float(cb[0]), float(cb[1]), float(cb[2]), float(cb[3])))

            if not chars:
                continue
            text = ''.join(c for c, *_ in chars)
            records.append({'bbox': bbox, 'text': text, 'chars': chars})
    return records

def _find_pdf_line_chars(line, line_text, raw_line_records):
    """按 bbox + 文本近似匹配当前行对应的 rawdict 字符列表。"""
    bbox = line.get('bbox')
    if not bbox or len(bbox) < 4:
        return None
    lx0, ly0, lx1, ly1 = map(float, bbox[:4])

    target = re.sub(r'\s+', '', line_text or '')
    best = None
    best_score = None
    for rec in raw_line_records:
        rb = rec.get('bbox')
        if not rb or len(rb) < 4:
            continue
        rx0, ry0, rx1, ry1 = map(float, rb[:4])
        # 先按 y 接近筛选
        if abs(ry0 - ly0) > 3.2 and abs(ry1 - ly1) > 3.2:
            continue

        rtext = re.sub(r'\s+', '', rec.get('text', ''))
        if not rtext:
            continue
        # 文本需有显著重叠
        if target and not (target in rtext or rtext in target):
            if target[:8] not in rtext and rtext[:8] not in target:
                continue

        score = abs(rx0 - lx0) + abs(ry0 - ly0) + abs(rx1 - lx1) + abs(ry1 - ly1)
        if best_score is None or score < best_score:
            best_score = score
            best = rec

    return best.get('chars') if best else None

def _extract_underlined_segments_from_pdf_line(line, line_text, page_words, underline_rects, line_chars=None):
    """按线条几何位置，从一行中提取被下划线标注的词片段。"""
    if not underline_rects:
        return []

    bbox = line.get('bbox')
    if not bbox or len(bbox) < 4:
        return []

    # lx0/ly0/lx1/ly1: 当前文本行的边界框坐标（左上/右下）。
    lx0, ly0, lx1, ly1 = map(float, bbox[:4])

    # 先尝试字符级匹配（最精确，适配中文无空格文本）。
    if line_chars:
        char_segments = []
        current = []
        for c, cx0, cy0, cx1, cy1 in line_chars:
            if not c.strip():
                if current:
                    seg = ''.join(current).strip(' ，,。；;：:、()（）[]【】')
                    if seg:
                        char_segments.append(seg)
                    current = []
                continue

            # under: 当前字符是否被某条下划线覆盖。
            under = False
            cw = max(cx1 - cx0, 0.1)
            for ux0, uy0, ux1, uy1 in underline_rects:
                if ux1 < cx0 or ux0 > cx1:
                    continue
                overlap = min(cx1, ux1) - max(cx0, ux0)
                if overlap < max(1.0, cw * 0.2):
                    continue
                if (cy1 - 1.4) <= uy0 <= (cy1 + 4.2):
                    under = True
                    break

            if under:
                current.append(c)
            else:
                if current:
                    seg = ''.join(current).strip(' ，,。；;：:、()（）[]【】')
                    if seg:
                        char_segments.append(seg)
                    current = []

        if current:
            seg = ''.join(current).strip(' ，,。；;：:、()（）[]【】')
            if seg:
                char_segments.append(seg)

        if char_segments:
            dedup = []
            seen_c = set()
            for s in char_segments:
                s2 = re.sub(r'\s+', '', s)
                if not s2 or s2 in seen_c:
                    continue
                seen_c.add(s2)
                dedup.append(s2)
            if dedup:
                out = [
                    s for s in dedup
                    if 2 <= len(s) <= 28 and not re.search(r'[，,。；;：:!?！？]', s)
                ]
                clean_line = re.sub(r'\s+', '', line_text or '')
                ratio = sum(len(s) for s in out) / max(len(clean_line), 1)
                if out and ratio <= 0.45 and len(out) <= 6:
                    return out

    # 取与该行 y 区间重叠的词（词级回退）。
    # candidates: 与当前行几何位置重叠的词级单元（词级回退用）。
    candidates = []
    for w in page_words:
        wx0, wy0, wx1, wy1, wtxt = float(w[0]), float(w[1]), float(w[2]), float(w[3]), str(w[4])
        if not wtxt.strip():
            continue
        if wy1 < ly0 - 2 or wy0 > ly1 + 2:
            continue
        if wx1 < lx0 - 2 or wx0 > lx1 + 2:
            continue
        candidates.append((wx0, wy0, wx1, wy1, wtxt))

    if not candidates:
        return []

    candidates.sort(key=lambda x: (x[1], x[0]))

    def is_underlined(word_box):
        wx0, wy0, wx1, wy1, _ = word_box
        ww = max(wx1 - wx0, 0.1)
        for ux0, uy0, ux1, uy1 in underline_rects:
            if ux1 < wx0 or ux0 > wx1:
                continue
            overlap = min(wx1, ux1) - max(wx0, ux0)
            if overlap < max(3.0, ww * 0.28):
                continue
            # 下划线 y 应位于词底部附近。
            if (wy1 - 1.5) <= uy0 <= (wy1 + 4.5):
                return True
        return False

    segments = []
    current = []
    last_x1 = None
    for wb in candidates:
        wx0, _, wx1, _, wtxt = wb
        under = is_underlined(wb)
        if under:
            if current and last_x1 is not None and wx0 - last_x1 > 8:
                seg = ''.join(current).strip(' ，,。；;：:、')
                if seg:
                    segments.append(seg)
                current = []
            current.append(wtxt)
            last_x1 = wx1
        else:
            if current:
                seg = ''.join(current).strip(' ，,。；;：:、')
                if seg:
                    segments.append(seg)
                current = []
                last_x1 = None

    if current:
        seg = ''.join(current).strip(' ，,。；;：:、')
        if seg:
            segments.append(seg)

    # 去重并过滤极短片段
    out = []
    seen = set()
    for s in segments:
        s2 = re.sub(r'\s+', '', s)
        if len(s2) < 2:
            continue
        if s2 in seen:
            continue
        seen.add(s2)
        out.append(s2)

    # 词级提取为空时，回退到几何字符映射（适配中文连续文本无空格的 PDF）。
    if not out:
        raw_line = (line_text or '').strip()
        if raw_line and (lx1 - lx0) > 1:
            n = len(raw_line)
            for ux0, uy0, ux1, uy1 in underline_rects:
                if ux1 < lx0 or ux0 > lx1:
                    continue
                # 下划线 y 需在该行底部附近。
                if not ((ly1 - 1.5) <= uy0 <= (ly1 + 4.5)):
                    continue

                rx0 = max(lx0, ux0)
                rx1 = min(lx1, ux1)
                if rx1 - rx0 < 8:
                    continue

                i0 = int(round((rx0 - lx0) / (lx1 - lx0) * n))
                i1 = int(round((rx1 - lx0) / (lx1 - lx0) * n))
                i0 = max(0, min(n - 1, i0))
                i1 = max(i0 + 1, min(n, i1))

                # 边界补全：下划线宽度可能略短，适度向两侧补齐中文词尾。
                punct = set(' ，,。；;：:!?！？、（）()[]【】')
                while i0 > 0 and raw_line[i0 - 1] not in punct and not raw_line[i0 - 1].isspace() and (i1 - i0) < 14:
                    i0 -= 1
                    # 左侧最多补 1 个字符，避免过扩
                    break
                while i1 < n and raw_line[i1] not in punct and not raw_line[i1].isspace() and (i1 - i0) < 14:
                    i1 += 1
                    # 右侧最多补 2 个字符，优先补齐词尾
                    if (i1 - i0) >= 2:
                        break

                seg = raw_line[i0:i1].strip(' ，,。；;：:、()（）[]【】')
                seg = re.sub(r'\s+', '', seg)
                # 去掉误吸附的短前缀（如“国、全面从严治党” -> “全面从严治党”）。
                while '、' in seg:
                    head, tail = seg.split('、', 1)
                    if len(head) <= 2 and len(tail) >= 2:
                        seg = tail
                    else:
                        break
                if seg:
                    out.append(seg)

        # 几何回退去重
        dedup = []
        seen2 = set()
        for s in out:
            if s in seen2:
                continue
            seen2.add(s)
            dedup.append(s)
        out = dedup

    # 高置信过滤：仅保留“短语级”片段，避免把整句误当答案。
    out = [
        s for s in out
        if 2 <= len(s) <= 28 and not re.search(r'[，,。；;：:!?！？]', s)
    ]
    if not out:
        return []

    clean_line = re.sub(r'\s+', '', line_text or '')
    ratio = sum(len(s) for s in out) / max(len(clean_line), 1)
    if ratio > 0.35 or len(out) > 4:
        return []

    return out

def _build_blank_question_from_line(line_text, segments):
    """将样式答案片段替换为填空位，生成“多空填空题”。"""
    question = line_text
    answers = []

    for idx, seg in enumerate(segments, 1):
        placeholder = f'（{idx}）______'
        if seg in question:
            question = question.replace(seg, placeholder, 1)
            answers.append(f'（{idx}）{seg}')

    if not answers:
        return None, None
    return question, _clean_answer_text('；'.join(answers))

def _extract_pdf_text(filepath):
    """提取 PDF 文本；若可获取字体信息，则尝试把样式答案转成“答案：”行。"""
    # 优先 PyMuPDF（支持 span 样式）
    try:
        import fitz  # type: ignore

        lines = []
        styled_pairs = []
        blank_pairs = []
        paragraph_pairs = []
        stream_pairs = []
        all_line_items = []
        with fitz.open(filepath) as doc:
            for page in doc:
                data = page.get_text('dict')
                raw_data = page.get_text('rawdict')
                raw_line_records = _build_pdf_raw_line_records(raw_data)
                page_words = page.get_text('words')
                underline_rects = _extract_pdf_underline_rects(page)
                for block in data.get('blocks', []):
                    if block.get('type') != 0:
                        continue

                    block_lines = []
                    block_segments = []
                    for line in block.get('lines', []):
                        spans = line.get('spans', [])
                        line_text = ''.join(s.get('text', '') for s in spans).strip()
                        if not line_text:
                            continue

                        block_lines.append(line_text)
                        line_segments = _extract_styled_segments_from_spans(spans)
                        line_chars = _find_pdf_line_chars(line, line_text, raw_line_records)
                        if not line_segments:
                            line_segments = _extract_underlined_segments_from_pdf_line(
                                line, line_text, page_words, underline_rects, line_chars
                            )
                        block_segments.extend(line_segments)
                        all_line_items.append((line_text, line_segments))

                        styled = ''.join(s.get('text', '') for s in spans if _pdf_span_is_styled(s)).strip()
                        lines.append(line_text)
                        if styled and styled != line_text and (QUESTION_START_RE.match(line_text) or _is_blank_question(line_text)):
                            lines.append(f'答案：{styled}')

                        # 优先生成“多空填空”题：将行内红字/样式片段替换为空位。
                        segments = _extract_styled_segments_from_spans(spans)
                        if not segments:
                            segments = _extract_underlined_segments_from_pdf_line(
                                line, line_text, page_words, underline_rects, line_chars
                            )
                        if (
                            segments and
                            len(line_text) <= 220 and
                            not _has_option_structure([line_text]) and
                            ('，' in line_text or '。' in line_text or ',' in line_text)
                        ):
                            q_line, a_line = _build_blank_question_from_line(line_text, segments)
                            if q_line and a_line and q_line != line_text:
                                blank_pairs.append((q_line, a_line))

                        # 兼容“红色/下划线大字即答案”这类PDF样式标注。
                        styled_ratio = len(styled) / max(len(line_text), 1)
                        if (
                            styled and styled != line_text and
                            0.05 <= styled_ratio <= 0.7 and
                            len(line_text) <= 180 and len(styled) <= 80 and
                            not _has_option_structure([line_text]) and
                            ('，' in line_text or '。' in line_text or ',' in line_text)
                        ):
                            styled_pairs.append((line_text, styled))

                    # 逐段多空：将一个文本块中的样式片段统一生成一道填空题。
                    if block_lines and block_segments:
                        block_text = _normalize_pdf_sentence(' '.join(block_lines))
                        block_segments = _dedupe_keep_order(block_segments)
                        if (
                            20 <= len(block_text) <= 800 and
                            1 <= len(block_segments) <= 16 and
                            not _has_option_structure([block_text]) and
                            ('，' in block_text or '。' in block_text or ',' in block_text)
                        ):
                            q_line, a_line = _build_blank_question_from_line(block_text, block_segments)
                            if q_line and a_line and q_line != block_text:
                                paragraph_pairs.append((q_line, a_line))

        # 跨页聚合：按阅读顺序拼接行，构造“整句/整段多空”题。
        buf_lines = []
        buf_segments = []
        for line_text, line_segments in all_line_items:
            buf_lines.append(line_text)
            buf_segments.extend(line_segments)

            paragraph_text = _normalize_pdf_sentence(' '.join(buf_lines))
            line_tail = line_text.rstrip()
            should_flush = line_tail.endswith(('。', '！', '？', '.', '!', '?')) or len(paragraph_text) >= 560
            if not should_flush:
                continue

            segs = _dedupe_keep_order(buf_segments)
            segs = _merge_split_segments(paragraph_text, segs)
            if (
                segs and
                30 <= len(paragraph_text) <= 1500 and
                1 <= len(segs) <= 24 and
                not _has_option_structure([paragraph_text]) and
                ('，' in paragraph_text or '。' in paragraph_text or ',' in paragraph_text)
            ):
                q_line, a_line = _build_blank_question_from_line(paragraph_text, segs)
                if q_line and a_line and q_line != paragraph_text:
                    stream_pairs.append((q_line, a_line))

            buf_lines = []
            buf_segments = []

        if buf_lines:
            paragraph_text = _normalize_pdf_sentence(' '.join(buf_lines))
            segs = _dedupe_keep_order(buf_segments)
            segs = _merge_split_segments(paragraph_text, segs)
            if (
                segs and
                30 <= len(paragraph_text) <= 1500 and
                1 <= len(segs) <= 24 and
                not _has_option_structure([paragraph_text]) and
                ('，' in paragraph_text or '。' in paragraph_text or ',' in paragraph_text)
            ):
                q_line, a_line = _build_blank_question_from_line(paragraph_text, segs)
                if q_line and a_line and q_line != paragraph_text:
                    stream_pairs.append((q_line, a_line))

        # 优先跨行整段多空，其次逐段/逐行，最后回退到样式答案对。
        generated_pairs = stream_pairs if stream_pairs else (paragraph_pairs if paragraph_pairs else (blank_pairs if blank_pairs else styled_pairs))

        if generated_pairs:
            lines.append('')
            for i, (q_text, ans_text) in enumerate(generated_pairs, 1):
                lines.append(f'{i}. {q_text}')
                lines.append(f'答案：{ans_text}')
        return _normalize_extracted_text('\n'.join(lines).strip())
    except Exception:
        pass

    # 回退到 pypdf / PyPDF2（仅文本）
    try:
        try:
            from pypdf import PdfReader  # type: ignore
        except Exception:
            from PyPDF2 import PdfReader  # type: ignore

        reader = PdfReader(filepath)
        pages = []
        for p in reader.pages:
            pages.append((p.extract_text() or '').strip())
        return _normalize_extracted_text('\n'.join(x for x in pages if x))
    except Exception:
        return None

def extract_text_by_filetype(filepath):
    """按扩展名读取题库文本，支持 txt/pdf/doc/docx。"""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.txt':
        return _read_text_file(filepath)

    if ext == '.docx':
        text = _extract_docx_text_with_style(filepath)
        if text:
            return text
        return _extract_docx_text_fallback(filepath)

    if ext == '.doc':
        text = _extract_doc_text_windows(filepath)
        if text:
            return text
        raise ValueError('读取 .doc 失败：请安装 Microsoft Word（COM）或先另存为 .docx/.txt 再导入。')

    if ext == '.pdf':
        text = _extract_pdf_text(filepath)
        if text:
            return text
        raise ValueError('读取 PDF 失败：请先安装 PyMuPDF 或 pypdf，或先转为 .txt/.docx。')

    raise ValueError(f'暂不支持的文件类型：{ext}')

def _normalize_answer_text(answer_text):
    """标准化答案文本：去空白、转大写、全角字母转半角。"""
    text = re.sub(r'\s+', '', answer_text or '').upper()
    # 全角字母转半角，兼容 DOC/PDF 中的答案格式。
    text = text.translate(str.maketrans('ＡＢＣＤＥＦＧＨ', 'ABCDEFGH'))
    return text

def _clean_answer_text(answer_text):
    """清理答案末尾孤立数字噪声（如换行后的 1/2/3）。"""
    t = (answer_text or '').strip()
    t = re.sub(r'\n\s*\d+\s*$', '', t)
    t = re.sub(r'[；;]\s*\d+\s*$', '', t)
    return t.strip()

def _is_answer_index_noise_line(text):
    """判断是否为答案区常见的序号噪声行（如 1. / 23 / （4））。"""
    t = str(text or '').strip()
    if not t:
        return False
    return bool(re.fullmatch(r'[（(]?\d{1,4}[)）]?[.。]?', t))

def _extract_choice_answer(answer_text, options):
    """尽量从答案文本中提取选项字母（支持 A/B/C、AB、A,C 等格式）。"""
    normalized = _normalize_answer_text(answer_text)
    letters = re.findall(r'[A-H]', normalized)
    if letters:
        return sorted(set(letters), key=letters.index)

    if len(options) == 2:
        # 判断题常见写法：答案为“正确/错误、对/错、是/否”
        true_words = ('正确', '对', '是')
        false_words = ('错误', '错', '否', '不正确')
        for k, v in options.items():
            opt_text = v.strip()
            if any(w in answer_text for w in true_words) and any(w in opt_text for w in true_words):
                return [k]
            if any(w in answer_text for w in false_words) and any(w in opt_text for w in false_words):
                return [k]

    # 答案直接写了选项文本时，尝试反查
    clean_answer = (answer_text or '').strip()
    if clean_answer:
        for k, v in options.items():
            opt = v.strip()
            if opt and (opt in clean_answer or clean_answer in opt):
                return [k]

    return []

def _is_judge_options(options):
    """根据选项文本判断是否为判断题（不依赖答案行）。"""
    if set(options.keys()) != {'A', 'B'}:
        return False

    a_text = options.get('A', '')
    b_text = options.get('B', '')
    judge_words = ('正确', '错误', '对', '错', '是', '否')
    return any(w in a_text for w in judge_words) and any(w in b_text for w in judge_words)

def _is_blank_question(question_text):
    """判断题干是否包含填空题常见占位符。"""
    blank_patterns = [
        r'[_＿﹍]{2,}',
        r'（\s*）',
        r'\(\s*\)',
        r'【\s*】',
        r'\[\s*\]',
        r'（\s*[_＿﹍\s]+\s*）',
        r'\(\s*[_＿﹍\s]+\s*\)',
        r'填空'
    ]
    return any(re.search(p, question_text) for p in blank_patterns)

def _looks_like_choice_answer_text(text):
    """判断文本是否像客观题答案串（字母或对错词）。"""
    normalized = _normalize_answer_text(text)
    cleaned = re.sub(r'[\s,，、;；/\\]+', '', normalized)
    if re.fullmatch(r'[A-H]+', cleaned or ''):
        return True
    return any(k in text for k in ('正确', '错误', '对', '错', '是', '否'))

def _split_content_and_answer(lines):
    """将题块行拆分为题干行与答案行，支持无前缀答案在下一行的极端格式。"""
    if len(lines) < 2:
        return None, None

    # 1) 优先匹配显式答案行（支持答案内容在下一行或多行）
    for i, line in enumerate(lines):
        m = ANSWER_LABEL_RE.match(line)
        if not m:
            continue

        head = m.group(1).strip()
        tail_lines = [l.strip() for l in lines[i + 1:] if l.strip()]
        content_lines = list(lines[:i])

        # 常规：答案与“答案：”同一行。
        if head:
            answer_lines = [head]
            if tail_lines:
                answer_lines.extend(tail_lines)
            answer_text = '\n'.join(answer_lines).strip()
            if content_lines and answer_text:
                return content_lines, answer_text
            continue

        # 兼容 PDF："答案："后常混入题号噪声与正文续行，避免把续行吞成答案。
        if not tail_lines:
            continue

        answer_idx = None
        for j in range(len(tail_lines) - 1, -1, -1):
            if _looks_like_answer_token(tail_lines[j]):
                answer_idx = j
                break

        if answer_idx is not None:
            for j, l in enumerate(tail_lines):
                if j >= answer_idx:
                    break
                if _is_answer_index_noise_line(l):
                    continue
                if _looks_like_question_start_line(l):
                    continue
                content_lines.append(l)

            if content_lines:
                return content_lines, tail_lines[answer_idx]
            continue

        # 未识别到明确答案：将非噪声续行并回题干，不返回答案。
        for l in tail_lines:
            if _is_answer_index_noise_line(l):
                continue
            if _looks_like_question_start_line(l):
                continue
            content_lines.append(l)

        if content_lines:
            return content_lines, ''

    # 2) 无答案前缀时的启发式拆分
    # 2.1 客观题常见：最后一行为标准答案 token（不是选项行）
    last_line = lines[-1].strip()
    content_lines = lines[:-1]
    if content_lines and _looks_like_answer_token(last_line):
        return content_lines, last_line

    # 2.2 填空/简答常见：最后一行为答案（尤其存在挖空标记）
    joined = ' '.join(lines)
    if (
        _is_blank_question(joined)
        and not _has_option_structure(content_lines)
        and not _has_option_structure([last_line])
    ):
        return content_lines, last_line

    # 2.25 简答常见：首行是问题（问号结尾），后续行为答案。
    has_option_in_tail = _has_option_structure(lines[1:])
    if len(lines) >= 2 and ('？' in lines[0] or '?' in lines[0]) and not has_option_in_tail:
        return [lines[0]], '\n'.join(lines[1:]).strip()

    # 2.3 默认视为未识别答案，不强拆最后一行，避免把选项误判为答案。
    return lines, ''

def _parse_questions_loose_qa(text):
    """宽松回退：从连续文本中抽取“问题行 + 后续答案行”的简答/填空。"""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if not lines:
        return []

    def is_question_line(line):
        if len(line) > 180:
            return False
        if '？' in line or '?' in line:
            return True

        # 一些资料用“概念：/特点：/内容：”而不写问号，也应当按提问行处理。
        if line.endswith(('：', ':')):
            if any(h in line for h in QA_PROMPT_HINTS):
                return True
        return False

    questions = []
    i = 0
    while i < len(lines):
        if not is_question_line(lines[i]):
            i += 1
            continue

        q_line = re.sub(r'^\s*(?:\d{1,4}|[一二三四五六七八九十百零]+)[.、．\)]\s*', '', lines[i]).strip()
        i += 1
        ans = []
        while i < len(lines) and not is_question_line(lines[i]):
            # 避免把章节标题吞进答案
            if QUESTION_START_RE.match(lines[i]) and len(lines[i]) < 25 and '？' not in lines[i] and '?' not in lines[i]:
                break
            ans.append(lines[i])
            i += 1

        answer_text = '\n'.join(ans).strip()
        if not q_line:
            continue

        # 无答案的问题行通常是目录/分节提示，过滤掉以降低噪声。
        if not answer_text:
            continue

        questions.append({
            'id': 0,
            'text': q_line,
            'options': {},
            'answer': answer_text,
            'type': 'blank' if _is_blank_question(q_line) else 'short'
        })

    for idx, q in enumerate(questions):
        q['id'] = idx + 1

    return questions

def _parse_numbered_qa_blocks(text):
    """解析“每题一行问题 + 后续多行答案”的编号简答题格式。"""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if not lines:
        return []

    blocks = []
    current = []
    for line in lines:
        if _looks_like_question_start_line(line):
            if current:
                blocks.append(current)
            current = [line]
        else:
            if current:
                current.append(line)
    if current:
        blocks.append(current)

    questions = []
    for block in blocks:
        if len(block) < 2:
            # 只有题干没有答案，不纳入可刷题集合。
            continue

        question_line = block[0]
        if _has_option_structure(block):
            continue

        question_text = QUESTION_START_RE.sub('', question_line).strip()
        answer_text = _clean_answer_text('\n'.join(block[1:]).strip())

        if not question_text or not answer_text:
            continue

        questions.append({
            'id': 0,
            'text': question_text,
            'options': {},
            'answer': answer_text,
            'type': 'blank' if _is_blank_question(question_text) else 'short'
        })

    for idx, q in enumerate(questions, 1):
        q['id'] = idx
    return questions

def _convert_text_to_judge_question(text):
    """将末尾带“对/错”标记的文本转为判断题。"""
    t = re.sub(r'\s+', ' ', str(text or '')).strip()
    if not t:
        return None

    m = re.search(r'(对|错|正确|错误)\s*$', t)
    if not m:
        return None

    token = m.group(1)
    stem = t[:m.start()].strip(' ：:;；，,。')
    if len(stem) < 6:
        return None

    answer = ['A'] if token in ('对', '正确') else ['B']
    return {
        'text': stem,
        'options': {'A': '正确', 'B': '错误'},
        'answer': answer,
        'type': 'judge'
    }

def _extract_judge_token(text):
    """从文本中提取判断题答案标记。"""
    t = re.sub(r'\s+', '', str(text or ''))
    if not t:
        return ''

    if re.fullmatch(r'(对|错|正确|错误)', t):
        return '对' if t in ('对', '正确') else '错'

    m = re.search(r'(对|错|正确|错误)', t)
    if not m:
        return ''
    return '对' if m.group(1) in ('对', '正确') else '错'

def _clean_pdf_judge_text_noise(text):
    """清理 PDF 判断题题干中的编号噪声与拼接残片。"""
    t = re.sub(r'\s+', ' ', str(text or '')).strip()
    if not t:
        return t

    # 去掉夹在中文语句中的孤立编号（常见于“答案：xx”误混入题干）。
    t = re.sub(r'(?<=[\u4e00-\u9fff，,、。；;：:])\s+\d{1,3}\s+(?=[\u4e00-\u9fff])', ' ', t)
    # 去掉“对 54 在...”这类下一题起始拼接。
    t = re.sub(r'\s*(?:对|错|正确|错误)\s+\d{1,3}\s+[\u4e00-\u9fff].*$', '', t)
    t = re.sub(r'\s{2,}', ' ', t).strip(' .．。 ，,；;')
    return t

def _split_compound_placeholder_judge_questions(questions):
    """把“（1）___ ... 对 （2）___ ... 错 ...”的大块题拆成多个判断题。"""
    marker_re = re.compile(r'（\s*\d+\s*）\s*[_＿﹍]{2,}')
    out = []

    for q in questions:
        q_type = q.get('type')
        if q_type not in ('blank', 'short') or q.get('options'):
            out.append(q)
            continue

        text = re.sub(r'\s+', ' ', str(q.get('text', '') or '')).strip()
        markers = list(marker_re.finditer(text))
        if len(markers) < 3:
            out.append(q)
            continue

        split_items = []
        for i, m in enumerate(markers):
            start = m.end()
            end = markers[i + 1].start() if i + 1 < len(markers) else len(text)
            seg = text[start:end].strip(' ；;，,。')
            if not seg:
                continue

            ans_m = re.search(r'(对|错|正确|错误)\s*$', seg)
            if not ans_m:
                continue
            token = ans_m.group(1)
            stem = seg[:ans_m.start()].strip(' ：:;；，,。')
            stem = _clean_pdf_judge_text_noise(stem)
            if len(stem) < 8:
                continue

            split_items.append({
                'id': 0,
                'text': stem,
                'options': {'A': '正确', 'B': '错误'},
                'answer': ['A'] if token in ('对', '正确') else ['B'],
                'type': 'judge'
            })

        if len(split_items) >= 3:
            out.extend(split_items)
        else:
            out.append(q)

    return out

def _postprocess_pdf_to_judge(questions):
    """PDF后处理：把“无选项+末尾对错”的短答转为判断题。"""
    for q in questions:
        if q.get('options'):
            continue
        if q.get('type') != 'short':
            continue

        q_text = re.sub(r'\s+', ' ', str(q.get('text', '') or '')).strip()
        q_answer = re.sub(r'\s+', ' ', str(q.get('answer', '') or '')).strip()

        # 先用题干本身判断，避免把“答案区数字”拼进题干。
        converted = _convert_text_to_judge_question(q_text)
        if not converted and q_answer and not _looks_like_answer_token(q_answer) and len(q_answer) >= 8:
            converted = _convert_text_to_judge_question((q_text + ' ' + q_answer).strip())

        if not converted:
            token = _extract_judge_token(q.get('answer', ''))
            if not token:
                continue
            converted = {
                'text': q_text,
                'options': {'A': '正确', 'B': '错误'},
                'answer': ['A'] if token == '对' else ['B'],
                'type': 'judge'
            }

        q['text'] = _clean_pdf_judge_text_noise(converted['text'])
        q['options'] = converted['options']
        q['answer'] = converted['answer']
        q['type'] = converted['type']

def _coerce_no_option_questions_to_judge(questions):
    """强制模式：把无选项题统一转为判断题（答案尽量从文本中提取）。"""
    out = []
    for q in questions:
        copied = dict(q)
        if copied.get('options'):
            out.append(copied)
            continue

        q_text = re.sub(r'\s+', ' ', str(copied.get('text', '') or '')).strip()
        q_answer = re.sub(r'\s+', ' ', str(copied.get('answer', '') or '')).strip()

        converted = _convert_text_to_judge_question(q_text)
        if not converted and q_answer and not _looks_like_answer_token(q_answer) and len(q_answer) >= 8:
            converted = _convert_text_to_judge_question((q_text + ' ' + q_answer).strip())

        if converted:
            copied['text'] = _clean_pdf_judge_text_noise(converted['text'])
            copied['options'] = converted['options']
            copied['answer'] = converted['answer']
            copied['type'] = converted['type']
        else:
            token = _extract_judge_token(copied.get('answer', ''))
            copied['text'] = _clean_pdf_judge_text_noise(copied.get('text', ''))
            copied['options'] = {'A': '正确', 'B': '错误'}
            copied['answer'] = ['A'] if token == '对' else (['B'] if token == '错' else [])
            copied['type'] = 'judge'

        out.append(copied)

    for i, q in enumerate(out, 1):
        q['id'] = i
    return out

def parse_questions(filepath):
    """解析题库文件，提取所有题目、选项、正确答案（支持 txt/pdf/doc/docx）。"""
    # 先做按文件类型的高置信专用解析，再走通用分块流程并做回填/后处理。
    ext = os.path.splitext(filepath)[1].lower()
    preferred_docx_questions = []

    # docx 优先尝试“红色选项直读”，避免回退文本解析误判答案。
    if ext == '.docx':
        numbered_choice_questions = _parse_docx_numbered_choice_questions(filepath)
        if numbered_choice_questions:
            answered = sum(1 for q in numbered_choice_questions if q.get('options') and q.get('answer'))
            # 自动编号试卷一般是纯选择题，命中到稳定规模时优先采用。
            if len(numbered_choice_questions) >= 10 and answered >= max(5, len(numbered_choice_questions) // 3):
                preferred_docx_questions = numbered_choice_questions

        if not preferred_docx_questions:
            red_docx_questions = _parse_docx_questions_with_red(filepath)
            styled_blank_questions = _parse_docx_styled_blank_questions(filepath)
            if red_docx_questions:
                objective_with_answer = sum(
                    1 for q in red_docx_questions
                    if q.get('options') and q.get('answer')
                )
                # 当检测到足量客观题且有有效答案时，优先保留红字解析结果。
                if len(red_docx_questions) >= 10 and objective_with_answer >= max(5, len(red_docx_questions) // 6):
                    preferred_docx_questions = _merge_docx_blank_questions(red_docx_questions, styled_blank_questions)

    text = extract_text_by_filetype(filepath)
    text = _normalize_extracted_text(text)

    questions = []

    # 优先按题号分块，兼容题目后答案、题目后选项+答案等不同结构。
    lines = text.split('\n')
    # current_block: 当前正在累积的单题文本块。
    current_block = []
    # blocks: [(题块文本, section_hint)]，section_hint 来源于“单选题/多选题”等分区标题。
    blocks = []
    has_question_boundary = False
    current_section = None

    for line in lines:
        sec = _detect_section_heading(line.strip())
        if sec:
            current_section = sec
            continue

        if _looks_like_question_start_line(line.strip()):
            has_question_boundary = True
            if current_block and any(l.strip() for l in current_block):
                blocks.append(('\n'.join(current_block).strip(), current_section))
                current_block = []
        current_block.append(line)

    if current_block and any(l.strip() for l in current_block):
        blocks.append(('\n'.join(current_block).strip(), current_section))

    # 如果按题号分块失败，则退化为按答案行分块。
    if not has_question_boundary:
        current_block = []
        blocks = []
        current_section = None
        for line in lines:
            sec = _detect_section_heading(line.strip())
            if sec:
                current_section = sec
                continue

            current_block.append(line)
            if ANSWER_LABEL_RE.match(line.strip()):
                blocks.append(('\n'.join(current_block).strip(), current_section))
                current_block = []

    for block, section_hint in blocks:
        q = parse_single_block(block)
        if q:
            if section_hint:
                q['section_hint'] = section_hint
            questions.append(q)

    # answer_keys: 文末答案区解析结果，结构为 {single/multi/judge/generic: {题号: 答案}}。
    answer_keys = _extract_answer_keys_from_text(text)
    _fill_answers_from_answer_keys(questions, answer_keys)

    # 低命中时回退到“编号简答”与“宽松问答”抽取，覆盖 txt/doc/docx/pdf 简答场景。
    if ext in ('.txt', '.pdf', '.doc', '.docx'):
        numbered_qa_questions = _parse_numbered_qa_blocks(text)
        loose_questions = _parse_questions_loose_qa(text)
        fallback_questions = numbered_qa_questions if len(numbered_qa_questions) >= len(loose_questions) else loose_questions

        if not questions:
            questions = fallback_questions
        else:
            # 若常规解析仅得到极少主观题，而宽松解析数量显著更多，
            # 则判定发生了“整篇并题”，切换为宽松结果。
            all_subjective = all(not q.get('options') for q in questions)
            if all_subjective and len(questions) <= 3 and len(fallback_questions) >= len(questions) + 5:
                questions = fallback_questions

    if ext == '.docx' and preferred_docx_questions:
        questions = _merge_docx_preferred_with_short_questions(preferred_docx_questions, questions)

    if ext == '.docx':
        styled_blank_questions = _parse_docx_styled_blank_questions(filepath)
        blank_count = sum(1 for q in questions if q.get('type') == 'blank')
        if styled_blank_questions and (blank_count == 0 or len(questions) <= 5):
            questions = _merge_docx_blank_questions(questions, styled_blank_questions)

    if ext == '.pdf':
        # PDF 可能因断行导致“题干+下一行正文”被误拆成短答，这里先做通用回并。
        for q in questions:
            if q.get('type') != 'short' or q.get('options'):
                continue

            q_text = str(q.get('text', '') or '')
            q_ans = str(q.get('answer', '') or '')

            # 若“答案”看起来像正文续句而非标准答案，先回并到题干。
            ans_is_continuation = (
                q_ans and
                not _looks_like_choice_answer_text(q_ans) and
                not q_ans.strip().startswith('参考答案') and
                len(re.sub(r'\s+', '', q_ans)) >= 8
            )

            merged_text = (q_text + q_ans).strip() if ans_is_continuation else q_text.strip()
            merged_text = re.sub(r'\s+', '', merged_text)
            if not merged_text:
                continue

            if ans_is_continuation:
                # 仅回并正文，不强行当作已知填空。
                q['text'] = merged_text
                q['answer'] = ''

        _postprocess_pdf_to_judge(questions)
        questions = _split_compound_placeholder_judge_questions(questions)
        for q in questions:
            if q.get('type') == 'judge':
                q['text'] = _clean_pdf_judge_text_noise(q.get('text', ''))

    # 重新编号
    for q in questions:
        if q.get('type') == 'short':
            ans = str(q.get('answer', '') or '')
            if re.search(r'（\s*\d+\s*）\s*[^；;\n]+', ans):
                q['type'] = 'blank'
                q['text'] = _mask_blank_question_text(q.get('text', ''), ans)

    for i, q in enumerate(questions):
        q['id'] = i + 1

    return questions

def parse_single_block(block):
    """解析单个题块，兼容客观题与主观题。"""
    # 单题块解析顺序：拆题干/答案 -> 识别选项 -> 判题型 -> 组装统一结构。
    lines = [l.strip() for l in block.strip().split('\n') if l.strip()]
    if not lines:
        return None

    # source_no: 原文题号（若能识别），用于后续答案回填与定位。
    source_no = None
    m_no = re.match(r'^\s*(\d{1,4})\s*[.、．\)]', lines[0])
    if m_no:
        source_no = int(m_no.group(1))

    content_lines, answer_text = _split_content_and_answer(lines)
    if not content_lines:
        return None
    answer_text = _clean_answer_text(answer_text)

    # 纠错：如果“答案”本身呈现选项结构，说明拆分过早，应回并到题干继续按客观题解析。
    if answer_text and _has_option_structure([answer_text]) and not _has_option_structure(content_lines):
        content_lines = content_lines + [answer_text]
        answer_text = ''

    # 将所有内容拼成一个长字符串，用换行分隔
    full_text = '\n'.join(content_lines)

    # 预处理：把紧挨在一起的选项拆开（如 "A.正确B.错误" → "A.正确\nB.错误"）
    # 匹配 非选项字母/非空白 后紧跟 选项字母+标点
    full_text = re.sub(r'(?<=[^\n])([A-H])[.、．,，]\s*', r'\n\1. ', full_text)
    # 清理可能产生的多余换行
    full_text = re.sub(r'\n{3,}', '\n\n', full_text)

    # 策略：先用正则拆出所有选项位置
    # 匹配 A. A、A，A) A）A: 等写法
    option_pattern = re.compile(r'(?:^|\n|[ \t]+)([A-H])[.、．,，\)）:]\s*', re.MULTILINE)

    # options: 选项映射，如 {'A': 'xxx', 'B': 'yyy'}。
    options = {}
    option_positions = []
    # format_marked_letters: 从“勾选符号样式”推断出的正确选项字母。
    format_marked_letters = []

    for m in option_pattern.finditer(full_text):
        option_positions.append((m.group(1), m.start(), m.end()))

    if not option_positions:
        # 尝试更宽松的匹配：A.xxx 紧贴前文
        option_pattern2 = re.compile(r'([A-H])[.、．,，\)）:]\s*')
        for m in option_pattern2.finditer(full_text):
            option_positions.append((m.group(1), m.start(), m.end()))

    if option_positions:
        # 去重：同一个字母只保留第一次出现（在题目文本之后的）
        # seen: 去重同字母重复匹配，避免 A 选项被多次识别。
        seen = set()
        unique_positions = []
        for letter, start, end in option_positions:
            if letter not in seen:
                seen.add(letter)
                unique_positions.append((letter, start, end))
        option_positions = unique_positions

        # 题目文本 = 第一个选项之前的内容
        first_opt_start = option_positions[0][1]
        question_text = full_text[:first_opt_start].strip()

        # 提取每个选项的文本
        for i, (letter, start, end) in enumerate(option_positions):
            if i + 1 < len(option_positions):
                next_start = option_positions[i + 1][1]
                opt_text = full_text[end:next_start].strip()
            else:
                opt_text = full_text[end:].strip()

            # 识别“选项字母前”的勾选标记（如：B、）。
            pre_context = full_text[max(0, start - 3):start]
            if _option_contains_answer_marker(pre_context):
                format_marked_letters.append(letter)

            # 清理选项文本
            opt_text = opt_text.replace('\n', ' ').strip()
            # 去掉末尾的标点
            opt_text = opt_text.rstrip('。.，,')
            options[letter] = _clean_option_text(opt_text)
    else:
        # 主观题：没有选项，整个内容就是题干
        question_text = full_text.strip()

    # 清理题目文本
    question_text = question_text.replace('\n', ' ')
    question_text = re.sub(r'[\uf0b7\u2022]+', ' ', question_text)
    # 移除开头的编号
    question_text = re.sub(r'^[\d]+[.、．\s]+', '', question_text)
    # 移除 【单选题】【多选题】 等标签
    question_text = re.sub(r'【[^】]*】\s*', '', question_text)
    # 移除章节标题如 "二.认真认识党史国史..." "三. 一起读马克思"
    question_text = re.sub(r'^[一二三四五六七八九十]+[.、．]\s*[\u4e00-\u9fff]+\s*', '', question_text)
    question_text = question_text.strip()

    if not question_text:
        return None

    # 过滤噪声：既无选项也无答案时，不作为可答题目。
    if not options and not (answer_text or '').strip():
        return None

    # 确定题目类型
    if options:
        # 关键修复：只要有选项，就按客观题判型，不再降级成简答/填空。
        correct_answers = _extract_choice_answer(answer_text or '', options)
        if not correct_answers:
            # PDF/文档中常见：正确项用“勾选符号”直接标在选项行。
            if format_marked_letters:
                correct_answers = sorted(set(format_marked_letters), key=format_marked_letters.index)
            else:
                correct_answers = _extract_choice_answer_from_option_format(options)
        if _is_judge_options(options):
            q_type = 'judge'
        elif len(correct_answers) > 1:
            q_type = 'multi'
        else:
            q_type = 'single'

        # 兜底：若答案未识别出字母，尝试从原始答案文本里提取；仍失败则保留为空列表。
        answer_value = correct_answers
        if not answer_value:
            normalized = _normalize_answer_text(answer_text or '')
            fallback_letters = re.findall(r'[A-H]', normalized)
            if fallback_letters:
                answer_value = sorted(set(fallback_letters), key=fallback_letters.index)
    else:
        q_type = 'blank' if _is_blank_question(question_text) else 'short'
        answer_value = answer_text or ''

    return {
        'id': 0,
        'text': question_text,
        'options': options,
        'answer': answer_value,
        'type': q_type,
        'source_no': source_no
    }

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
题库刷题工具
功能：
1. 手动选择题库文件，提取所有题目和答案，重新编号
2. 图形化界面，随机抽题、选择作答、显示正确答案
3. 记录错误次数到 error_record.json，支持断点续记
4. 根据错误次数和错误率进行加权随机抽题
"""
# Version 3.1.3

import tkinter as tk
from tkinter import messagebox, font as tkfont, filedialog, ttk
import re
import json
import os
import random
import math
import ctypes
import tempfile
import zipfile
import xml.etree.ElementTree as ET

# ============ 题目解析模块 ============

ANSWER_LABEL_RE = re.compile(r'^(?:正确答案|答案|参考答案|标准答案|【答案】|\[答案\]|答|参考解答)\s*[：:]?\s*(.*)\s*$')
QUESTION_START_RE = re.compile(r'^\s*(?:\d{1,4}(?:[、．\)]|[.](?!\d))|[一二三四五六七八九十百零]+[.、．\)])\s*')
OPTION_PREFIX_RE = re.compile(r'^\s*([A-HＡ-Ｈ])[.、．,，\)）:：]\s*')
OPTION_TOKEN_RE = re.compile(r'([A-HＡ-Ｈ])[.、．,，\)）:：]\s*')

# 宽松问答解析中，用于识别“无问号但明显是提问提示语”的行。
QA_PROMPT_HINTS = (
    '什么是', '哪些', '哪几', '哪一', '如何', '为什么', '为何', '是否',
    '概念', '特点', '内容', '内涵', '作用', '意义', '影响', '要求',
    '区别', '分类', '趋势', '原则', '方法', '地位', '历程', '阶段',
    '名称', '含义', '实质', '时间', '对象', '标准', '功能', '战略',
)


def _normalize_extracted_text(text):
    """规范化不同来源文本，提升题块边界与选项识别稳定性。"""
    if not text:
        return ''

    t = text.replace('\r\n', '\n').replace('\r', '\n')
    t = t.replace('\u3000', ' ')
    # 标准答案标签前强制换行
    t = re.sub(r'(?<!\n)(正确答案|参考答案|标准答案|答案|答)[：:]', r'\n\1：', t)
    # 题号前强制换行，兼容 OCR/PDF 黏连；避免把小数(如4.5)误判成题号。
    t = re.sub(r'(?<!\n)(\d{1,4}(?:[、．\)]|[.](?!\d)))', r'\n\1', t)
    t = re.sub(r'\n{3,}', '\n\n', t)
    return t.strip()


def _looks_like_option_line(text):
    if not text:
        return False
    if OPTION_PREFIX_RE.match(text):
        return True
    markers = OPTION_TOKEN_RE.findall(text)
    return len(markers) >= 2


def _has_option_structure(lines):
    """跨多行判断是否包含客观题选项结构。"""
    if not lines:
        return False

    prefixed = sum(1 for line in lines if OPTION_PREFIX_RE.match(line))
    if prefixed >= 2:
        return True

    joined = ' '.join(lines)
    markers = OPTION_TOKEN_RE.findall(joined)
    return len(markers) >= 2


def _looks_like_answer_token(text):
    """判断文本是否像标准答案（而不是选项行）。"""
    if not text:
        return False
    if _looks_like_option_line(text):
        return False

    normalized = _normalize_answer_text(text)
    cleaned = re.sub(r'[\s,，、;；/\\]+', '', normalized)
    if re.fullmatch(r'[A-H]+', cleaned or ''):
        return True

    answer_keywords = ('正确', '错误', '对', '错', '是', '否')
    return any(k == text.strip() for k in answer_keywords)


def _read_text_file(filepath):
    for enc in ('utf-8-sig', 'utf-8', 'utf-16', 'utf-16-le', 'utf-16-be', 'gb18030', 'gbk'):
        try:
            with open(filepath, 'r', encoding=enc) as f:
                return f.read()
        except Exception:
            continue
    raise ValueError(f'无法读取文本文件编码：{filepath}')


def _extract_docx_text_with_style(filepath):
    """优先用 python-docx 提取，尽量利用字体样式识别答案。"""
    try:
        from docx import Document
    except Exception:
        return None

    def run_is_styled(run):
        font = run.font
        color = None
        if font and font.color is not None and font.color.rgb is not None:
            color = str(font.color.rgb)
        return bool(
            run.bold or run.italic or run.underline or
            (font and font.highlight_color is not None) or
            color
        )

    lines = []
    doc = Document(filepath)

    def consume_para(para):
        raw = ''.join(r.text for r in para.runs).strip()
        if not raw:
            return
        styled = ''.join(r.text for r in para.runs if run_is_styled(r)).strip()
        lines.append(raw)

        # 若答案仅以不同字体标记，则注入标准答案行
        if styled and styled != raw and (QUESTION_START_RE.match(raw) or _is_blank_question(raw)):
            lines.append(f'答案：{styled}')

    for para in doc.paragraphs:
        consume_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    consume_para(para)

    return _normalize_extracted_text('\n'.join(lines).strip())


def _docx_run_is_red(run):
    """判断 docx run 是否为红色（兼容 RGB 与底层 XML 颜色值）。"""
    try:
        color = run.font.color if run.font else None
        if color and color.rgb:
            rgb = str(color.rgb).upper()
            if rgb.startswith('FF') and rgb[2:4] in ('00', '11', '22', '33') and rgb[4:6] in ('00', '11', '22', '33'):
                return True

        color_el = run._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
        if color_el is not None:
            val = (color_el.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') or '').upper().strip('#')
            if val in {'FF0000', 'C00000', 'CC0000', 'DD0000', 'EE0000'}:
                return True
    except Exception:
        return False
    return False


def _docx_run_is_nonblack(run):
    """判断 run 是否为非黑色字体（含主题色/显式颜色）。"""

    def is_near_black(rgb_hex):
        if not rgb_hex or len(rgb_hex) != 6:
            return False
        try:
            r = int(rgb_hex[0:2], 16)
            g = int(rgb_hex[2:4], 16)
            b = int(rgb_hex[4:6], 16)
        except Exception:
            return False

        # 深灰正文（如 333333）不应被视作“强调色”。
        return max(r, g, b) <= 70 and (max(r, g, b) - min(r, g, b)) <= 12

    try:
        color = run.font.color if run.font else None
        if color is None:
            return False

        if color.rgb is not None:
            rgb = str(color.rgb).upper()
            rgb = rgb.replace('#', '')
            if len(rgb) == 8:
                rgb = rgb[-6:]
            if rgb and rgb not in {'000000', '00000000'} and not is_near_black(rgb):
                return True
    except Exception:
        return False
    return False


def _docx_has_nondefault_shading(element):
    """检测 XML 节点上的底纹/背景色是否为非默认值。"""
    if element is None:
        return False

    try:
        shd_nodes = element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
    except Exception:
        shd_nodes = []

    for shd in shd_nodes:
        fill = (shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill') or '').upper()
        color = (shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color') or '').upper()
        val = (shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') or '').lower()

        # 空值、auto、白色视作默认背景；其余颜色或样式视作有背景标记。
        if fill and fill not in {'AUTO', 'FFFFFF', '00000000'}:
            return True
        if color and color not in {'AUTO', 'FFFFFF', '000000', '00000000'}:
            return True
        if val and val not in {'clear', 'nil'}:
            return True

    return False


def _docx_run_is_emphasis(run):
    """用于填空提取：下划线/加粗/非黑色/高亮均视为强调。"""
    if _docx_run_is_nonblack(run):
        return True
    if bool(getattr(run, 'underline', False)):
        return True
    if bool(getattr(run, 'bold', False)):
        return True
    try:
        if run.font and run.font.highlight_color is not None:
            return True
    except Exception:
        pass

    # 支持底层 XML 的背景色/底纹（某些文档把答案标在背景色而不是高亮里）。
    try:
        if _docx_has_nondefault_shading(run._element):
            return True
        parent = getattr(run._element, 'getparent', lambda: None)()
        if parent is not None and _docx_has_nondefault_shading(parent):
            return True
    except Exception:
        pass

    return False


def _parse_docx_styled_blank_questions(filepath):
    """从 docx 样式中提取填空题：强调样式文本视为答案片段。"""
    try:
        from docx import Document
    except Exception:
        return []

    try:
        doc = Document(filepath)
    except Exception:
        return []

    def iter_paragraphs():
        for p in doc.paragraphs:
            yield p
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    questions = []
    seen_texts = set()

    for para in iter_paragraphs():
        raw = para.text.strip()
        if not raw:
            continue
        if _looks_like_option_line(raw):
            continue

        full_text = ''.join((r.text or '') for r in para.runs).strip()
        if not full_text:
            full_text = raw

        if len(full_text) < 12 or len(full_text) > 260:
            continue
        if not any(ch in full_text for ch in ('，', '。', '；', ';', '：', ':', ',')):
            continue
        if _has_option_structure([full_text]):
            continue

        # 提取连续强调片段
        segments = []
        current = []
        for run in para.runs:
            t = (run.text or '')
            if not t.strip():
                continue
            if _docx_run_is_emphasis(run):
                current.append(t)
            else:
                if current:
                    seg = ''.join(current).strip(' ，,。；;：:、()（）[]【】')
                    if seg:
                        segments.append(seg)
                    current = []
        if current:
            seg = ''.join(current).strip(' ，,。；;：:、()（）[]【】')
            if seg:
                segments.append(seg)

        # 去重并过滤噪声
        cleaned_segments = []
        seen_seg = set()
        for seg in segments:
            seg = re.sub(r'\s+', '', seg)
            if not seg or seg in seen_seg:
                continue
            if len(seg) > 40:
                continue
            if re.fullmatch(r'[\W_]+', seg):
                continue
            if not re.search(r'[A-Za-z0-9\u4e00-\u9fff]', seg):
                continue
            seen_seg.add(seg)
            cleaned_segments.append(seg)

        if not cleaned_segments:
            continue

        # 过高强调占比多为整句加粗标题，不当作填空。
        ratio = sum(len(s) for s in cleaned_segments) / max(len(full_text), 1)
        if ratio < 0.03 or ratio > 0.55:
            continue
        if len(cleaned_segments) > 10:
            continue

        question_text = full_text
        answers = []
        replaced = False
        for idx, seg in enumerate(cleaned_segments, 1):
            placeholder = f'（{idx}）______'
            if seg in question_text:
                question_text = question_text.replace(seg, placeholder, 1)
                answers.append(f'（{idx}）{seg}')
                replaced = True

        if not replaced:
            continue

        question_text = QUESTION_START_RE.sub('', question_text).strip()
        if question_text in seen_texts:
            continue
        seen_texts.add(question_text)

        questions.append({
            'id': 0,
            'text': question_text,
            'options': {},
            'answer': _clean_answer_text('；'.join(answers)),
            'type': 'blank'
        })

    for idx, q in enumerate(questions, 1):
        q['id'] = idx
    return questions


def _merge_docx_blank_questions(base_questions, blank_questions):
    """将样式填空题并入解析结果，避免重复题干。"""
    if not blank_questions:
        return base_questions

    def normalize_for_match(text):
        t = (text or '')
        t = re.sub(r'（\s*\d+\s*）\s*[_＿﹍]+', ' ', t)
        t = re.sub(r'[_＿﹍]+', ' ', t)
        t = re.sub(r'[\s，,。；;：:、（）()\[\]【】]+', '', t)
        return t

    blank_anchors = []
    for b in blank_questions:
        txt = b.get('text', '')
        parts = re.split(r'（\s*\d+\s*）\s*[_＿﹍]+', txt)
        anchors = [normalize_for_match(p) for p in parts if normalize_for_match(p)]
        blank_anchors.append(anchors)

    filtered_base = []
    for q in base_questions:
        if q.get('options'):
            filtered_base.append(q)
            continue

        q_text_norm = normalize_for_match(q.get('text', ''))
        should_drop = False
        if q.get('type') in ('short', 'blank') and q_text_norm:
            for anchors in blank_anchors:
                if not anchors:
                    continue
                hit_count = sum(1 for a in anchors if len(a) >= 6 and a in q_text_norm)
                long_hit = any(len(a) >= 12 and a in q_text_norm for a in anchors)
                if long_hit or hit_count >= 2:
                    should_drop = True
                    break

        if not should_drop:
            filtered_base.append(q)

    merged = list(filtered_base)
    existing = set()
    for q in merged:
        t = re.sub(r'\s+', '', q.get('text', ''))
        if t:
            existing.add(t)

    for b in blank_questions:
        t = re.sub(r'\s+', '', b.get('text', ''))
        if not t or t in existing:
            continue
        merged.append(dict(b))
        existing.add(t)

    for idx, q in enumerate(merged, 1):
        q['id'] = idx
    return merged


def _mask_blank_question_text(question_text, answer_text):
    """兜底：若填空题题干未挖空，则根据答案片段自动替换为空位。"""
    text = (question_text or '').strip()
    if not text:
        return text
    if _is_blank_question(text):
        return text

    ans = (answer_text or '').strip()
    if not ans:
        return text

    # 优先匹配“（1）答案；（2）答案”格式。
    pairs = re.findall(r'（\s*(\d+)\s*）\s*([^；;\n]+)', ans)
    if pairs:
        out = text
        replaced = False
        for idx, (_, seg) in enumerate(pairs, 1):
            seg = seg.strip()
            if seg and seg in out:
                out = out.replace(seg, f'（{idx}）______', 1)
                replaced = True
        if replaced:
            return out

    # 次级策略：分号切分答案片段后替换。
    chunks = [c.strip() for c in re.split(r'[；;]+', ans) if c.strip()]
    out = text
    replaced = 0
    for i, chunk in enumerate(chunks, 1):
        chunk = re.sub(r'^（\s*\d+\s*）', '', chunk).strip()
        if chunk and chunk in out:
            out = out.replace(chunk, f'（{i}）______', 1)
            replaced += 1

    return out if replaced > 0 else text


def build_parse_candidates(filepath):
    """构建题库解析候选结果，供预览时切换比对。"""
    candidates = []
    ext = os.path.splitext(filepath)[1].lower()

    auto_questions = parse_questions(filepath)
    candidates.append(('自动识别（推荐）', auto_questions, '综合策略自动选择并合并结果'))

    if ext == '.docx':
        red_questions = _parse_docx_questions_with_red(filepath)
        if red_questions:
            candidates.append(('仅红色选项识别', red_questions, '按红色字体识别客观题答案'))

        styled_blanks = _parse_docx_styled_blank_questions(filepath)
        if styled_blanks:
            candidates.append(('仅样式填空识别', styled_blanks, '将下划线/非黑色/加粗视为填空答案'))

        if red_questions and styled_blanks:
            merged = _merge_docx_blank_questions(red_questions, styled_blanks)
            candidates.append(('红色客观 + 样式填空', merged, '客观题用红色，填空题用样式提取'))

    unique = []
    seen = set()
    for name, qs, desc in candidates:
        signature = (len(qs), tuple(sorted((q.get('type', ''), len(q.get('options', {}))) for q in qs[:50])))
        key = (name, signature)
        if key in seen:
            continue
        seen.add(key)
        unique.append((name, qs, desc))
    return unique


def _docx_extract_option_segments(para_text):
    """从段落文本中提取选项段：(letter, start, end, option_text)。"""
    segments = []
    marker_re = re.compile(r'([A-HＡ-Ｈ])[.、．,，\)）:：]\s*')
    markers = list(marker_re.finditer(para_text))
    if not markers:
        return segments

    for i, m in enumerate(markers):
        letter = m.group(1).translate(str.maketrans('ＡＢＣＤＥＦＧＨ', 'ABCDEFGH'))
        start = m.start()
        content_start = m.end()
        content_end = markers[i + 1].start() if i + 1 < len(markers) else len(para_text)
        option_text = para_text[content_start:content_end].strip().rstrip('。.，,；;')
        segments.append((letter, start, content_end, option_text))
    return segments


def _parse_docx_questions_with_red(filepath):
    """直接解析 docx：按红色选项识别正确答案，避免文本回退误判。"""
    try:
        from docx import Document
    except Exception:
        return []

    try:
        doc = Document(filepath)
    except Exception:
        return []

    def iter_paragraphs():
        for p in doc.paragraphs:
            yield p
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    questions = []
    current = None

    def finalize_current():
        nonlocal current
        if not current:
            return

        question_text = current['text']
        question_text = re.sub(r'^[\d]+[.、．\s]+', '', question_text)
        question_text = re.sub(r'【[^】]*】\s*', '', question_text)
        question_text = question_text.strip()

        if not question_text or not current['options']:
            current = None
            return

        red_answers = sorted(current['red_options'])
        if not red_answers and current['explicit_answer']:
            red_answers = _extract_choice_answer(current['explicit_answer'], current['options'])

        if _is_judge_options(current['options']):
            q_type = 'judge'
        elif len(red_answers) > 1:
            q_type = 'multi'
        else:
            q_type = 'single'

        questions.append({
            'id': 0,
            'text': question_text,
            'options': current['options'],
            'answer': red_answers,
            'type': q_type
        })
        current = None

    for para in iter_paragraphs():
        para_text = para.text.strip()
        if not para_text:
            continue

        if QUESTION_START_RE.match(para_text):
            finalize_current()
            current = {
                'text': para_text,
                'options': {},
                'red_options': set(),
                'explicit_answer': ''
            }
            continue

        if current is None:
            continue

        ans_match = ANSWER_LABEL_RE.match(para_text)
        if ans_match:
            current['explicit_answer'] = ans_match.group(1).strip()
            continue

        full_text = ''.join((r.text or '') for r in para.runs)
        if not full_text.strip():
            full_text = para_text

        segments = _docx_extract_option_segments(full_text)
        if not segments:
            # 续行题干
            if not current['options']:
                current['text'] = (current['text'] + ' ' + para_text).strip()
            continue

        # 构建字符级红色掩码，按“选项段”判断哪一项为红色。
        red_mask = []
        for run in para.runs:
            txt = run.text or ''
            if not txt:
                continue
            red_mask.extend([_docx_run_is_red(run)] * len(txt))
        if len(red_mask) < len(full_text):
            red_mask.extend([False] * (len(full_text) - len(red_mask)))

        for letter, start, end, option_text in segments:
            if option_text:
                current['options'][letter] = option_text

            # 仅检测本选项的有效区间：去掉尾部空白，避免“下个选项前红色空格”误伤当前选项。
            detect_end = end
            while detect_end > start and full_text[detect_end - 1].isspace():
                detect_end -= 1

            # 该选项有效区间任一字符为红色，即视为正确选项。
            if detect_end > start and any(red_mask[start:detect_end]):
                current['red_options'].add(letter)

    finalize_current()

    for idx, q in enumerate(questions, 1):
        q['id'] = idx
    return questions


def _extract_docx_text_fallback(filepath):
    """不依赖第三方库的 docx 文本提取回退方案。"""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    with zipfile.ZipFile(filepath, 'r') as zf:
        xml_data = zf.read('word/document.xml')

    root = ET.fromstring(xml_data)
    lines = []
    for p in root.findall('.//w:p', ns):
        texts = []
        for t in p.findall('.//w:t', ns):
            texts.append(t.text or '')
        line = ''.join(texts).strip()
        if line:
            lines.append(line)
    return _normalize_extracted_text('\n'.join(lines).strip())


def _extract_doc_text_windows(filepath):
    """使用 Windows Word COM 将 .doc/.docx 转成 txt（若本机可用）。"""
    try:
        import win32com.client  # type: ignore
    except Exception:
        return None

    temp_txt = None
    word = None
    doc = None
    try:
        fd, temp_txt = tempfile.mkstemp(suffix='.txt')
        os.close(fd)

        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(filepath))
        # 7 = wdFormatUnicodeText
        doc.SaveAs(os.path.abspath(temp_txt), FileFormat=7)
        doc.Close(False)
        doc = None
        word.Quit()
        word = None
        return _normalize_extracted_text(_read_text_file(temp_txt))
    except Exception:
        return None
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        if temp_txt and os.path.exists(temp_txt):
            try:
                os.remove(temp_txt)
            except Exception:
                pass


def _pdf_span_is_styled(span):
    font_name = str(span.get('font', '')).lower()
    color = span.get('color', 0)
    flags = int(span.get('flags', 0))
    if any(k in font_name for k in ('bold', 'black', 'heavy', 'demi', 'medium')):
        return True
    if color not in (0, None):
        return True
    # 低位标志中某些位常用于强调样式，作为启发式
    if flags & 2 or flags & 16:
        return True
    return False


def _extract_styled_segments_from_spans(spans):
    """从PDF行内span提取连续样式片段（红字/下划线/粗体等）。"""
    def style_signature(span):
        return (
            str(span.get('font', '')),
            round(float(span.get('size', 0) or 0), 2),
            int(span.get('flags', 0) or 0),
            int(span.get('color', 0) or 0)
        )

    valid_spans = [s for s in spans if (s.get('text', '') or '').strip()]
    if not valid_spans:
        return []

    # 计算本行“主样式”：按文本长度加权出现最多的样式。
    sig_weight = {}
    for s in valid_spans:
        sig = style_signature(s)
        sig_weight[sig] = sig_weight.get(sig, 0) + len((s.get('text', '') or '').strip())
    dominant_sig = max(sig_weight.items(), key=lambda kv: kv[1])[0]

    def is_contrasted_styled(span):
        # 仅将“有强调样式且不同于主样式”的片段视为答案候选，降低整行误判。
        return _pdf_span_is_styled(span) and style_signature(span) != dominant_sig

    segments = []
    current = []

    for s in valid_spans:
        txt = (s.get('text', '') or '').strip()
        if not txt:
            continue

        if is_contrasted_styled(s):
            current.append(txt)
        else:
            if current:
                seg = ''.join(current).strip(' ，,。；;：:、')
                if seg:
                    segments.append(seg)
                current = []

    if current:
        seg = ''.join(current).strip(' ，,。；;：:、')
        if seg:
            segments.append(seg)

    # 若完全没有“差异样式”，则回退为空，避免把整行普通文本误当答案。
    if not segments:
        return []

    # 去重（保持顺序）
    seen = set()
    unique = []
    for seg in segments:
        if seg in seen:
            continue
        # 过滤明显噪声
        if len(seg) < 2 and not re.fullmatch(r'[\u4e00-\u9fff]', seg):
            continue
        if re.fullmatch(r'[\W_]+', seg):
            continue
        seen.add(seg)
        unique.append(seg)
    return unique


def _dedupe_keep_order(items):
    seen = set()
    out = []
    for x in items:
        if x in seen:
            continue
        seen.add(x)
        out.append(x)
    return out


def _normalize_pdf_sentence(text):
    t = text or ''
    # 修复中文文本因换行导致的断词空格
    t = re.sub(r'(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])', '', t)
    t = re.sub(r'\s{2,}', ' ', t)
    return t.strip()


def _merge_split_segments(paragraph_text, segments):
    """合并被换行拆开的相邻答案片段（如“习近平新”+“时代...”）。"""
    if not segments:
        return segments

    merged = []
    i = 0
    while i < len(segments):
        current = segments[i]
        j = i + 1
        while j < len(segments):
            candidate = current + segments[j]
            if candidate in paragraph_text:
                current = candidate
                j += 1
            else:
                break
        merged.append(current)
        i = j

    return _dedupe_keep_order(merged)


def _extract_pdf_underline_rects(page):
    """提取 PDF 页面中的细横线矩形（常见于下划线标注答案）。"""
    rects = []
    try:
        drawings = page.get_drawings()
    except Exception:
        return rects

    for d in drawings:
        for it in d.get('items', []):
            if not it or it[0] != 're':
                continue
            r = it[1]
            x0, y0, x1, y1 = float(r.x0), float(r.y0), float(r.x1), float(r.y1)
            w = max(0.0, x1 - x0)
            h = max(0.0, y1 - y0)

            # 下划线通常是“很薄但较长”的横向矩形。
            if w >= 12 and h <= 1.2:
                rects.append((x0, y0, x1, y1))
    return rects


def _build_pdf_raw_line_records(raw_data):
    """从 rawdict 构建行级字符记录，供下划线字符级匹配。"""
    records = []
    for block in raw_data.get('blocks', []):
        if block.get('type') != 0:
            continue
        for line in block.get('lines', []):
            bbox = line.get('bbox')
            chars = []
            for span in line.get('spans', []):
                for ch in span.get('chars', []):
                    c = ch.get('c', '')
                    cb = ch.get('bbox')
                    if not c or not cb or len(cb) < 4:
                        continue
                    chars.append((c, float(cb[0]), float(cb[1]), float(cb[2]), float(cb[3])))

            if not chars:
                continue
            text = ''.join(c for c, *_ in chars)
            records.append({'bbox': bbox, 'text': text, 'chars': chars})
    return records


def _find_pdf_line_chars(line, line_text, raw_line_records):
    """按 bbox + 文本近似匹配当前行对应的 rawdict 字符列表。"""
    bbox = line.get('bbox')
    if not bbox or len(bbox) < 4:
        return None
    lx0, ly0, lx1, ly1 = map(float, bbox[:4])

    target = re.sub(r'\s+', '', line_text or '')
    best = None
    best_score = None
    for rec in raw_line_records:
        rb = rec.get('bbox')
        if not rb or len(rb) < 4:
            continue
        rx0, ry0, rx1, ry1 = map(float, rb[:4])
        # 先按 y 接近筛选
        if abs(ry0 - ly0) > 3.2 and abs(ry1 - ly1) > 3.2:
            continue

        rtext = re.sub(r'\s+', '', rec.get('text', ''))
        if not rtext:
            continue
        # 文本需有显著重叠
        if target and not (target in rtext or rtext in target):
            if target[:8] not in rtext and rtext[:8] not in target:
                continue

        score = abs(rx0 - lx0) + abs(ry0 - ly0) + abs(rx1 - lx1) + abs(ry1 - ly1)
        if best_score is None or score < best_score:
            best_score = score
            best = rec

    return best.get('chars') if best else None


def _extract_underlined_segments_from_pdf_line(line, line_text, page_words, underline_rects, line_chars=None):
    """按线条几何位置，从一行中提取被下划线标注的词片段。"""
    if not underline_rects:
        return []

    bbox = line.get('bbox')
    if not bbox or len(bbox) < 4:
        return []

    lx0, ly0, lx1, ly1 = map(float, bbox[:4])

    # 先尝试字符级匹配（最精确，适配中文无空格文本）。
    if line_chars:
        char_segments = []
        current = []
        for c, cx0, cy0, cx1, cy1 in line_chars:
            if not c.strip():
                if current:
                    seg = ''.join(current).strip(' ，,。；;：:、()（）[]【】')
                    if seg:
                        char_segments.append(seg)
                    current = []
                continue

            under = False
            cw = max(cx1 - cx0, 0.1)
            for ux0, uy0, ux1, uy1 in underline_rects:
                if ux1 < cx0 or ux0 > cx1:
                    continue
                overlap = min(cx1, ux1) - max(cx0, ux0)
                if overlap < max(1.0, cw * 0.2):
                    continue
                if (cy1 - 1.4) <= uy0 <= (cy1 + 4.2):
                    under = True
                    break

            if under:
                current.append(c)
            else:
                if current:
                    seg = ''.join(current).strip(' ，,。；;：:、()（）[]【】')
                    if seg:
                        char_segments.append(seg)
                    current = []

        if current:
            seg = ''.join(current).strip(' ，,。；;：:、()（）[]【】')
            if seg:
                char_segments.append(seg)

        if char_segments:
            dedup = []
            seen_c = set()
            for s in char_segments:
                s2 = re.sub(r'\s+', '', s)
                if not s2 or s2 in seen_c:
                    continue
                seen_c.add(s2)
                dedup.append(s2)
            if dedup:
                out = [
                    s for s in dedup
                    if 2 <= len(s) <= 28 and not re.search(r'[，,。；;：:!?！？]', s)
                ]
                clean_line = re.sub(r'\s+', '', line_text or '')
                ratio = sum(len(s) for s in out) / max(len(clean_line), 1)
                if out and ratio <= 0.45 and len(out) <= 6:
                    return out

    # 取与该行 y 区间重叠的词（词级回退）。
    candidates = []
    for w in page_words:
        wx0, wy0, wx1, wy1, wtxt = float(w[0]), float(w[1]), float(w[2]), float(w[3]), str(w[4])
        if not wtxt.strip():
            continue
        if wy1 < ly0 - 2 or wy0 > ly1 + 2:
            continue
        if wx1 < lx0 - 2 or wx0 > lx1 + 2:
            continue
        candidates.append((wx0, wy0, wx1, wy1, wtxt))

    if not candidates:
        return []

    candidates.sort(key=lambda x: (x[1], x[0]))

    def is_underlined(word_box):
        wx0, wy0, wx1, wy1, _ = word_box
        ww = max(wx1 - wx0, 0.1)
        for ux0, uy0, ux1, uy1 in underline_rects:
            if ux1 < wx0 or ux0 > wx1:
                continue
            overlap = min(wx1, ux1) - max(wx0, ux0)
            if overlap < max(3.0, ww * 0.28):
                continue
            # 下划线 y 应位于词底部附近。
            if (wy1 - 1.5) <= uy0 <= (wy1 + 4.5):
                return True
        return False

    segments = []
    current = []
    last_x1 = None
    for wb in candidates:
        wx0, _, wx1, _, wtxt = wb
        under = is_underlined(wb)
        if under:
            if current and last_x1 is not None and wx0 - last_x1 > 8:
                seg = ''.join(current).strip(' ，,。；;：:、')
                if seg:
                    segments.append(seg)
                current = []
            current.append(wtxt)
            last_x1 = wx1
        else:
            if current:
                seg = ''.join(current).strip(' ，,。；;：:、')
                if seg:
                    segments.append(seg)
                current = []
                last_x1 = None

    if current:
        seg = ''.join(current).strip(' ，,。；;：:、')
        if seg:
            segments.append(seg)

    # 去重并过滤极短片段
    out = []
    seen = set()
    for s in segments:
        s2 = re.sub(r'\s+', '', s)
        if len(s2) < 2:
            continue
        if s2 in seen:
            continue
        seen.add(s2)
        out.append(s2)

    # 词级提取为空时，回退到几何字符映射（适配中文连续文本无空格的 PDF）。
    if not out:
        raw_line = (line_text or '').strip()
        if raw_line and (lx1 - lx0) > 1:
            n = len(raw_line)
            for ux0, uy0, ux1, uy1 in underline_rects:
                if ux1 < lx0 or ux0 > lx1:
                    continue
                # 下划线 y 需在该行底部附近。
                if not ((ly1 - 1.5) <= uy0 <= (ly1 + 4.5)):
                    continue

                rx0 = max(lx0, ux0)
                rx1 = min(lx1, ux1)
                if rx1 - rx0 < 8:
                    continue

                i0 = int(round((rx0 - lx0) / (lx1 - lx0) * n))
                i1 = int(round((rx1 - lx0) / (lx1 - lx0) * n))
                i0 = max(0, min(n - 1, i0))
                i1 = max(i0 + 1, min(n, i1))

                # 边界补全：下划线宽度可能略短，适度向两侧补齐中文词尾。
                punct = set(' ，,。；;：:!?！？、（）()[]【】')
                while i0 > 0 and raw_line[i0 - 1] not in punct and not raw_line[i0 - 1].isspace() and (i1 - i0) < 14:
                    i0 -= 1
                    # 左侧最多补 1 个字符，避免过扩
                    break
                while i1 < n and raw_line[i1] not in punct and not raw_line[i1].isspace() and (i1 - i0) < 14:
                    i1 += 1
                    # 右侧最多补 2 个字符，优先补齐词尾
                    if (i1 - i0) >= 2:
                        break

                seg = raw_line[i0:i1].strip(' ，,。；;：:、()（）[]【】')
                seg = re.sub(r'\s+', '', seg)
                # 去掉误吸附的短前缀（如“国、全面从严治党” -> “全面从严治党”）。
                while '、' in seg:
                    head, tail = seg.split('、', 1)
                    if len(head) <= 2 and len(tail) >= 2:
                        seg = tail
                    else:
                        break
                if seg:
                    out.append(seg)

        # 几何回退去重
        dedup = []
        seen2 = set()
        for s in out:
            if s in seen2:
                continue
            seen2.add(s)
            dedup.append(s)
        out = dedup

    # 高置信过滤：仅保留“短语级”片段，避免把整句误当答案。
    out = [
        s for s in out
        if 2 <= len(s) <= 28 and not re.search(r'[，,。；;：:!?！？]', s)
    ]
    if not out:
        return []

    clean_line = re.sub(r'\s+', '', line_text or '')
    ratio = sum(len(s) for s in out) / max(len(clean_line), 1)
    if ratio > 0.35 or len(out) > 4:
        return []

    return out


def _build_blank_question_from_line(line_text, segments):
    """将样式答案片段替换为填空位，生成“多空填空题”。"""
    question = line_text
    answers = []

    for idx, seg in enumerate(segments, 1):
        placeholder = f'（{idx}）______'
        if seg in question:
            question = question.replace(seg, placeholder, 1)
            answers.append(f'（{idx}）{seg}')

    if not answers:
        return None, None
    return question, _clean_answer_text('；'.join(answers))


def _extract_pdf_text(filepath):
    """提取 PDF 文本；若可获取字体信息，则尝试把样式答案转成“答案：”行。"""
    # 优先 PyMuPDF（支持 span 样式）
    try:
        import fitz  # type: ignore

        lines = []
        styled_pairs = []
        blank_pairs = []
        paragraph_pairs = []
        stream_pairs = []
        all_line_items = []
        with fitz.open(filepath) as doc:
            for page in doc:
                data = page.get_text('dict')
                raw_data = page.get_text('rawdict')
                raw_line_records = _build_pdf_raw_line_records(raw_data)
                page_words = page.get_text('words')
                underline_rects = _extract_pdf_underline_rects(page)
                for block in data.get('blocks', []):
                    if block.get('type') != 0:
                        continue

                    block_lines = []
                    block_segments = []
                    for line in block.get('lines', []):
                        spans = line.get('spans', [])
                        line_text = ''.join(s.get('text', '') for s in spans).strip()
                        if not line_text:
                            continue

                        block_lines.append(line_text)
                        line_segments = _extract_styled_segments_from_spans(spans)
                        line_chars = _find_pdf_line_chars(line, line_text, raw_line_records)
                        if not line_segments:
                            line_segments = _extract_underlined_segments_from_pdf_line(
                                line, line_text, page_words, underline_rects, line_chars
                            )
                        block_segments.extend(line_segments)
                        all_line_items.append((line_text, line_segments))

                        styled = ''.join(s.get('text', '') for s in spans if _pdf_span_is_styled(s)).strip()
                        lines.append(line_text)
                        if styled and styled != line_text and (QUESTION_START_RE.match(line_text) or _is_blank_question(line_text)):
                            lines.append(f'答案：{styled}')

                        # 优先生成“多空填空”题：将行内红字/样式片段替换为空位。
                        segments = _extract_styled_segments_from_spans(spans)
                        if not segments:
                            segments = _extract_underlined_segments_from_pdf_line(
                                line, line_text, page_words, underline_rects, line_chars
                            )
                        if (
                            segments and
                            len(line_text) <= 220 and
                            not _has_option_structure([line_text]) and
                            ('，' in line_text or '。' in line_text or ',' in line_text)
                        ):
                            q_line, a_line = _build_blank_question_from_line(line_text, segments)
                            if q_line and a_line and q_line != line_text:
                                blank_pairs.append((q_line, a_line))

                        # 兼容“红色/下划线大字即答案”这类PDF样式标注。
                        styled_ratio = len(styled) / max(len(line_text), 1)
                        if (
                            styled and styled != line_text and
                            0.05 <= styled_ratio <= 0.7 and
                            len(line_text) <= 180 and len(styled) <= 80 and
                            not _has_option_structure([line_text]) and
                            ('，' in line_text or '。' in line_text or ',' in line_text)
                        ):
                            styled_pairs.append((line_text, styled))

                    # 逐段多空：将一个文本块中的样式片段统一生成一道填空题。
                    if block_lines and block_segments:
                        block_text = _normalize_pdf_sentence(' '.join(block_lines))
                        block_segments = _dedupe_keep_order(block_segments)
                        if (
                            20 <= len(block_text) <= 800 and
                            1 <= len(block_segments) <= 16 and
                            not _has_option_structure([block_text]) and
                            ('，' in block_text or '。' in block_text or ',' in block_text)
                        ):
                            q_line, a_line = _build_blank_question_from_line(block_text, block_segments)
                            if q_line and a_line and q_line != block_text:
                                paragraph_pairs.append((q_line, a_line))

        # 跨页聚合：按阅读顺序拼接行，构造“整句/整段多空”题。
        buf_lines = []
        buf_segments = []
        for line_text, line_segments in all_line_items:
            buf_lines.append(line_text)
            buf_segments.extend(line_segments)

            paragraph_text = _normalize_pdf_sentence(' '.join(buf_lines))
            line_tail = line_text.rstrip()
            should_flush = line_tail.endswith(('。', '！', '？', '.', '!', '?')) or len(paragraph_text) >= 560
            if not should_flush:
                continue

            segs = _dedupe_keep_order(buf_segments)
            segs = _merge_split_segments(paragraph_text, segs)
            if (
                segs and
                30 <= len(paragraph_text) <= 1500 and
                1 <= len(segs) <= 24 and
                not _has_option_structure([paragraph_text]) and
                ('，' in paragraph_text or '。' in paragraph_text or ',' in paragraph_text)
            ):
                q_line, a_line = _build_blank_question_from_line(paragraph_text, segs)
                if q_line and a_line and q_line != paragraph_text:
                    stream_pairs.append((q_line, a_line))

            buf_lines = []
            buf_segments = []

        if buf_lines:
            paragraph_text = _normalize_pdf_sentence(' '.join(buf_lines))
            segs = _dedupe_keep_order(buf_segments)
            segs = _merge_split_segments(paragraph_text, segs)
            if (
                segs and
                30 <= len(paragraph_text) <= 1500 and
                1 <= len(segs) <= 24 and
                not _has_option_structure([paragraph_text]) and
                ('，' in paragraph_text or '。' in paragraph_text or ',' in paragraph_text)
            ):
                q_line, a_line = _build_blank_question_from_line(paragraph_text, segs)
                if q_line and a_line and q_line != paragraph_text:
                    stream_pairs.append((q_line, a_line))

        # 优先跨行整段多空，其次逐段/逐行，最后回退到样式答案对。
        generated_pairs = stream_pairs if stream_pairs else (paragraph_pairs if paragraph_pairs else (blank_pairs if blank_pairs else styled_pairs))

        if generated_pairs:
            lines.append('')
            for i, (q_text, ans_text) in enumerate(generated_pairs, 1):
                lines.append(f'{i}. {q_text}')
                lines.append(f'答案：{ans_text}')
        return _normalize_extracted_text('\n'.join(lines).strip())
    except Exception:
        pass

    # 回退到 pypdf / PyPDF2（仅文本）
    try:
        try:
            from pypdf import PdfReader  # type: ignore
        except Exception:
            from PyPDF2 import PdfReader  # type: ignore

        reader = PdfReader(filepath)
        pages = []
        for p in reader.pages:
            pages.append((p.extract_text() or '').strip())
        return _normalize_extracted_text('\n'.join(x for x in pages if x))
    except Exception:
        return None


def extract_text_by_filetype(filepath):
    """按扩展名读取题库文本，支持 txt/pdf/doc/docx。"""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.txt':
        return _read_text_file(filepath)

    if ext == '.docx':
        text = _extract_docx_text_with_style(filepath)
        if text:
            return text
        return _extract_docx_text_fallback(filepath)

    if ext == '.doc':
        text = _extract_doc_text_windows(filepath)
        if text:
            return text
        raise ValueError('读取 .doc 失败：请安装 Microsoft Word（COM）或先另存为 .docx/.txt 再导入。')

    if ext == '.pdf':
        text = _extract_pdf_text(filepath)
        if text:
            return text
        raise ValueError('读取 PDF 失败：请先安装 PyMuPDF 或 pypdf，或先转为 .txt/.docx。')

    raise ValueError(f'暂不支持的文件类型：{ext}')


def _normalize_answer_text(answer_text):
    text = re.sub(r'\s+', '', answer_text or '').upper()
    # 全角字母转半角，兼容 DOC/PDF 中的答案格式。
    text = text.translate(str.maketrans('ＡＢＣＤＥＦＧＨ', 'ABCDEFGH'))
    return text


def _clean_answer_text(answer_text):
    """清理答案末尾孤立数字噪声（如换行后的 1/2/3）。"""
    t = (answer_text or '').strip()
    t = re.sub(r'\n\s*\d+\s*$', '', t)
    t = re.sub(r'[；;]\s*\d+\s*$', '', t)
    return t.strip()


def _extract_choice_answer(answer_text, options):
    """尽量从答案文本中提取选项字母（支持 A/B/C、AB、A,C 等格式）。"""
    normalized = _normalize_answer_text(answer_text)
    letters = re.findall(r'[A-H]', normalized)
    if letters:
        return sorted(set(letters), key=letters.index)

    if len(options) == 2:
        # 判断题常见写法：答案为“正确/错误、对/错、是/否”
        true_words = ('正确', '对', '是')
        false_words = ('错误', '错', '否', '不正确')
        for k, v in options.items():
            opt_text = v.strip()
            if any(w in answer_text for w in true_words) and any(w in opt_text for w in true_words):
                return [k]
            if any(w in answer_text for w in false_words) and any(w in opt_text for w in false_words):
                return [k]

    # 答案直接写了选项文本时，尝试反查
    clean_answer = (answer_text or '').strip()
    if clean_answer:
        for k, v in options.items():
            opt = v.strip()
            if opt and (opt in clean_answer or clean_answer in opt):
                return [k]

    return []


def _is_judge_options(options):
    """根据选项文本判断是否为判断题（不依赖答案行）。"""
    if set(options.keys()) != {'A', 'B'}:
        return False

    a_text = options.get('A', '')
    b_text = options.get('B', '')
    judge_words = ('正确', '错误', '对', '错', '是', '否')
    return any(w in a_text for w in judge_words) and any(w in b_text for w in judge_words)


def _is_blank_question(question_text):
    blank_patterns = [
        r'[_＿﹍]{2,}',
        r'（\s*）',
        r'\(\s*\)',
        r'【\s*】',
        r'\[\s*\]',
        r'（\s*[_＿﹍\s]+\s*）',
        r'\(\s*[_＿﹍\s]+\s*\)',
        r'填空'
    ]
    return any(re.search(p, question_text) for p in blank_patterns)


def _looks_like_choice_answer_text(text):
    normalized = _normalize_answer_text(text)
    cleaned = re.sub(r'[\s,，、;；/\\]+', '', normalized)
    if re.fullmatch(r'[A-H]+', cleaned or ''):
        return True
    return any(k in text for k in ('正确', '错误', '对', '错', '是', '否'))


def _split_content_and_answer(lines):
    """将题块行拆分为题干行与答案行，支持无前缀答案在下一行的极端格式。"""
    if len(lines) < 2:
        return None, None

    # 1) 优先匹配显式答案行（支持答案内容在下一行或多行）
    for i, line in enumerate(lines):
        m = ANSWER_LABEL_RE.match(line)
        if not m:
            continue

        head = m.group(1).strip()
        answer_lines = []
        if head:
            answer_lines.append(head)
        if i + 1 < len(lines):
            answer_lines.extend(l.strip() for l in lines[i + 1:] if l.strip())

        answer_text = '\n'.join(answer_lines).strip()
        content_lines = lines[:i]
        if content_lines and answer_text:
            return content_lines, answer_text

    # 2) 无答案前缀时的启发式拆分
    # 2.1 客观题常见：最后一行为标准答案 token（不是选项行）
    last_line = lines[-1].strip()
    content_lines = lines[:-1]
    if content_lines and _looks_like_answer_token(last_line):
        return content_lines, last_line

    # 2.2 填空/简答常见：最后一行为答案（尤其存在挖空标记）
    joined = ' '.join(lines)
    if (
        _is_blank_question(joined)
        and not _has_option_structure(content_lines)
        and not _has_option_structure([last_line])
    ):
        return content_lines, last_line

    # 2.25 简答常见：首行是问题（问号结尾），后续行为答案。
    has_option_in_tail = _has_option_structure(lines[1:])
    if len(lines) >= 2 and ('？' in lines[0] or '?' in lines[0]) and not has_option_in_tail:
        return [lines[0]], '\n'.join(lines[1:]).strip()

    # 2.3 默认视为未识别答案，不强拆最后一行，避免把选项误判为答案。
    return lines, ''


def _parse_questions_loose_qa(text):
    """宽松回退：从连续文本中抽取“问题行 + 后续答案行”的简答/填空。"""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if not lines:
        return []

    def is_question_line(line):
        if len(line) > 180:
            return False
        if '？' in line or '?' in line:
            return True

        # 一些资料用“概念：/特点：/内容：”而不写问号，也应当按提问行处理。
        if line.endswith(('：', ':')):
            if any(h in line for h in QA_PROMPT_HINTS):
                return True
        return False

    questions = []
    i = 0
    while i < len(lines):
        if not is_question_line(lines[i]):
            i += 1
            continue

        q_line = re.sub(r'^\s*(?:\d{1,4}|[一二三四五六七八九十百零]+)[.、．\)]\s*', '', lines[i]).strip()
        i += 1
        ans = []
        while i < len(lines) and not is_question_line(lines[i]):
            # 避免把章节标题吞进答案
            if QUESTION_START_RE.match(lines[i]) and len(lines[i]) < 25 and '？' not in lines[i] and '?' not in lines[i]:
                break
            ans.append(lines[i])
            i += 1

        answer_text = '\n'.join(ans).strip()
        if not q_line:
            continue

        # 无答案的问题行通常是目录/分节提示，过滤掉以降低噪声。
        if not answer_text:
            continue

        questions.append({
            'id': 0,
            'text': q_line,
            'options': {},
            'answer': answer_text,
            'type': 'blank' if _is_blank_question(q_line) else 'short'
        })

    for idx, q in enumerate(questions):
        q['id'] = idx + 1

    return questions


def _parse_numbered_qa_blocks(text):
    """解析“每题一行问题 + 后续多行答案”的编号简答题格式。"""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if not lines:
        return []

    blocks = []
    current = []
    for line in lines:
        if QUESTION_START_RE.match(line):
            if current:
                blocks.append(current)
            current = [line]
        else:
            if current:
                current.append(line)
    if current:
        blocks.append(current)

    questions = []
    for block in blocks:
        if len(block) < 2:
            # 只有题干没有答案，不纳入可刷题集合。
            continue

        question_line = block[0]
        if _has_option_structure(block):
            continue

        question_text = QUESTION_START_RE.sub('', question_line).strip()
        answer_text = _clean_answer_text('\n'.join(block[1:]).strip())

        if not question_text or not answer_text:
            continue

        questions.append({
            'id': 0,
            'text': question_text,
            'options': {},
            'answer': answer_text,
            'type': 'blank' if _is_blank_question(question_text) else 'short'
        })

    for idx, q in enumerate(questions, 1):
        q['id'] = idx
    return questions


def _apply_pdf_known_blank_patterns(text):
    """针对无样式信息的 PDF，按高置信固定表达做保底挖空。"""
    t = (text or '').strip()
    if not t:
        return None

    compact = re.sub(r'\s+', '', t)

    # 2020 秋初党常见表述：仅“伟大斗争”“伟大梦想”为填空位。
    pattern = re.compile(r'统揽\s*伟大斗争\s*[、，,]\s*伟大工程\s*[、，,]\s*伟大事业\s*[、，,]\s*伟大梦想')
    if pattern.search(compact):
        masked = pattern.sub('统揽（1）______、伟大工程、伟大事业、（2）______', compact)
        answer = '（1）伟大斗争；（2）伟大梦想'
        return masked, answer

    pattern2 = re.compile(r'可以延长预备期，但不能超过一年；不履行党员义务，不具备党员条件的，应当取消预备党员资格')
    if pattern2.search(compact):
        masked = pattern2.sub('可以延长预备期，但不能超过（1）______；不履行党员义务，不具备党员条件的，（2）______', compact)
        answer = '（1）一年；（2）应当取消预备党员资格'
        return masked, answer

    return None

def parse_questions(filepath):
    """解析题库文件，提取所有题目、选项、正确答案（支持 txt/pdf/doc/docx）。"""
    ext = os.path.splitext(filepath)[1].lower()

    # docx 优先尝试“红色选项直读”，避免回退文本解析误判答案。
    if ext == '.docx':
        red_docx_questions = _parse_docx_questions_with_red(filepath)
        styled_blank_questions = _parse_docx_styled_blank_questions(filepath)
        if red_docx_questions:
            objective_with_answer = sum(
                1 for q in red_docx_questions
                if q.get('options') and q.get('answer')
            )
            # 当检测到足量客观题且有有效答案时，直接采用红字解析结果。
            if len(red_docx_questions) >= 10 and objective_with_answer >= max(5, len(red_docx_questions) // 6):
                return _merge_docx_blank_questions(red_docx_questions, styled_blank_questions)

    text = extract_text_by_filetype(filepath)
    text = _normalize_extracted_text(text)

    questions = []

    # 优先按题号分块，兼容题目后答案、题目后选项+答案等不同结构。
    lines = text.split('\n')
    current_block = []
    blocks = []
    has_question_boundary = False

    for line in lines:
        if QUESTION_START_RE.match(line.strip()):
            has_question_boundary = True
            if current_block and any(l.strip() for l in current_block):
                blocks.append('\n'.join(current_block).strip())
                current_block = []
        current_block.append(line)

    if current_block and any(l.strip() for l in current_block):
        blocks.append('\n'.join(current_block).strip())

    # 如果按题号分块失败，则退化为按答案行分块。
    if not has_question_boundary:
        current_block = []
        blocks = []
        for line in lines:
            current_block.append(line)
            if ANSWER_LABEL_RE.match(line.strip()):
                blocks.append('\n'.join(current_block).strip())
                current_block = []

    for block in blocks:
        q = parse_single_block(block)
        if q:
            questions.append(q)

    # 低命中时回退到“编号简答”与“宽松问答”抽取，覆盖 txt/doc/docx/pdf 简答场景。
    if ext in ('.txt', '.pdf', '.doc', '.docx'):
        numbered_qa_questions = _parse_numbered_qa_blocks(text)
        loose_questions = _parse_questions_loose_qa(text)
        fallback_questions = numbered_qa_questions if len(numbered_qa_questions) >= len(loose_questions) else loose_questions

        if not questions:
            questions = fallback_questions
        else:
            # 若常规解析仅得到极少主观题，而宽松解析数量显著更多，
            # 则判定发生了“整篇并题”，切换为宽松结果。
            all_subjective = all(not q.get('options') for q in questions)
            if all_subjective and len(questions) <= 3 and len(fallback_questions) >= len(questions) + 5:
                questions = fallback_questions

    if ext == '.docx':
        styled_blank_questions = _parse_docx_styled_blank_questions(filepath)
        blank_count = sum(1 for q in questions if q.get('type') == 'blank')
        if styled_blank_questions and (blank_count == 0 or len(questions) <= 5):
            questions = _merge_docx_blank_questions(questions, styled_blank_questions)

    if ext == '.pdf':
        # PDF 可能因断行导致“题干+下一行正文”被误拆成短答，这里先合并再做保底挖空。
        for q in questions:
            if q.get('type') != 'short' or q.get('options'):
                continue

            q_text = str(q.get('text', '') or '')
            q_ans = str(q.get('answer', '') or '')

            # 若“答案”看起来像正文续句而非标准答案，先回并到题干。
            ans_is_continuation = (
                q_ans and
                not _looks_like_choice_answer_text(q_ans) and
                not q_ans.strip().startswith('参考答案') and
                len(re.sub(r'\s+', '', q_ans)) >= 8
            )

            merged_text = (q_text + q_ans).strip() if ans_is_continuation else q_text.strip()
            merged_text = re.sub(r'\s+', '', merged_text)
            if not merged_text:
                continue

            pattern_result = _apply_pdf_known_blank_patterns(merged_text)
            if pattern_result:
                masked, ans = pattern_result
                q['type'] = 'blank'
                q['text'] = masked
                q['answer'] = ans
            elif ans_is_continuation:
                # 仅回并正文，不强行当作已知填空。
                q['text'] = merged_text
                q['answer'] = ''

        # 对已识别为单空的“预备期满”题目做补空修正。
        for q in questions:
            if q.get('type') != 'blank':
                continue
            text_v = str(q.get('text', '') or '')
            if ('可以延长预备期' in text_v and '应当取消预备党员资格' in text_v and '（2）' not in text_v):
                q['text'] = text_v.replace('应当取消预备党员资格', '（2）______')
                ans_v = str(q.get('answer', '') or '')
                if '应当取消预备党员资格' not in ans_v:
                    if ans_v.strip():
                        q['answer'] = _clean_answer_text(ans_v + '；（2）应当取消预备党员资格')
                    else:
                        q['answer'] = '（2）应当取消预备党员资格'

    # 重新编号
    for q in questions:
        if q.get('type') == 'short':
            ans = str(q.get('answer', '') or '')
            if re.search(r'（\s*\d+\s*）\s*[^；;\n]+', ans):
                q['type'] = 'blank'
                q['text'] = _mask_blank_question_text(q.get('text', ''), ans)

    for i, q in enumerate(questions):
        q['id'] = i + 1

    return questions


def parse_single_block(block):
    """解析单个题块，兼容客观题与主观题。"""
    lines = [l.strip() for l in block.strip().split('\n') if l.strip()]
    if not lines:
        return None

    content_lines, answer_text = _split_content_and_answer(lines)
    if not content_lines:
        return None
    answer_text = _clean_answer_text(answer_text)

    # 纠错：如果“答案”本身呈现选项结构，说明拆分过早，应回并到题干继续按客观题解析。
    if answer_text and _has_option_structure([answer_text]) and not _has_option_structure(content_lines):
        content_lines = content_lines + [answer_text]
        answer_text = ''

    # 将所有内容拼成一个长字符串，用换行分隔
    full_text = '\n'.join(content_lines)

    # 预处理：把紧挨在一起的选项拆开（如 "A.正确B.错误" → "A.正确\nB.错误"）
    # 匹配 非选项字母/非空白 后紧跟 选项字母+标点
    full_text = re.sub(r'(?<=[^\n])([A-H])[.、．,，]\s*', r'\n\1. ', full_text)
    # 清理可能产生的多余换行
    full_text = re.sub(r'\n{3,}', '\n\n', full_text)

    # 策略：先用正则拆出所有选项位置
    # 匹配 A. A、A，A) A）A: 等写法
    option_pattern = re.compile(r'(?:^|\n|[ \t]+)([A-H])[.、．,，\)）:]\s*', re.MULTILINE)

    options = {}
    option_positions = []

    for m in option_pattern.finditer(full_text):
        option_positions.append((m.group(1), m.start(), m.end()))

    if not option_positions:
        # 尝试更宽松的匹配：A.xxx 紧贴前文
        option_pattern2 = re.compile(r'([A-H])[.、．,，\)）:]\s*')
        for m in option_pattern2.finditer(full_text):
            option_positions.append((m.group(1), m.start(), m.end()))

    if option_positions:
        # 去重：同一个字母只保留第一次出现（在题目文本之后的）
        seen = set()
        unique_positions = []
        for letter, start, end in option_positions:
            if letter not in seen:
                seen.add(letter)
                unique_positions.append((letter, start, end))
        option_positions = unique_positions

        # 题目文本 = 第一个选项之前的内容
        first_opt_start = option_positions[0][1]
        question_text = full_text[:first_opt_start].strip()

        # 提取每个选项的文本
        for i, (letter, start, end) in enumerate(option_positions):
            if i + 1 < len(option_positions):
                next_start = option_positions[i + 1][1]
                opt_text = full_text[end:next_start].strip()
            else:
                opt_text = full_text[end:].strip()
            # 清理选项文本
            opt_text = opt_text.replace('\n', ' ').strip()
            # 去掉末尾的标点
            opt_text = opt_text.rstrip('。.，,')
            options[letter] = opt_text
    else:
        # 主观题：没有选项，整个内容就是题干
        question_text = full_text.strip()

    # 清理题目文本
    question_text = question_text.replace('\n', ' ')
    # 移除开头的编号
    question_text = re.sub(r'^[\d]+[.、．\s]+', '', question_text)
    # 移除 【单选题】【多选题】 等标签
    question_text = re.sub(r'【[^】]*】\s*', '', question_text)
    # 移除章节标题如 "二.认真认识党史国史..." "三. 一起读马克思"
    question_text = re.sub(r'^[一二三四五六七八九十]+[.、．]\s*[\u4e00-\u9fff]+\s*', '', question_text)
    question_text = question_text.strip()

    if not question_text:
        return None

    # 过滤噪声：既无选项也无答案时，不作为可答题目。
    if not options and not (answer_text or '').strip():
        return None

    # 确定题目类型
    if options:
        # 关键修复：只要有选项，就按客观题判型，不再降级成简答/填空。
        correct_answers = _extract_choice_answer(answer_text or '', options)
        if _is_judge_options(options):
            q_type = 'judge'
        elif len(correct_answers) > 1:
            q_type = 'multi'
        else:
            q_type = 'single'

        # 兜底：若答案未识别出字母，尝试从原始答案文本里提取；仍失败则保留为空列表。
        answer_value = correct_answers
        if not answer_value:
            normalized = _normalize_answer_text(answer_text or '')
            fallback_letters = re.findall(r'[A-H]', normalized)
            if fallback_letters:
                answer_value = sorted(set(fallback_letters), key=fallback_letters.index)
    else:
        q_type = 'blank' if _is_blank_question(question_text) else 'short'
        answer_value = answer_text or ''

    return {
        'id': 0,
        'text': question_text,
        'options': options,
        'answer': answer_value,
        'type': q_type
    }


# ============ 错误记录模块 ============

RECORD_FILE = 'error_record.json'


def load_records():
    """加载历史错误记录"""
    if os.path.exists(RECORD_FILE):
        with open(RECORD_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}


def save_records(records):
    """保存错误记录"""
    with open(RECORD_FILE, 'w', encoding='utf-8') as f:
        json.dump(records, f, ensure_ascii=False, indent=2)


def get_record(records, qid):
    """获取某题的记录"""
    key = str(qid)
    if key in records:
        return records[key]
    return {'attempts': 0, 'errors': 0}


def update_record(records, qid, is_correct):
    """更新某题的记录"""
    key = str(qid)
    if key not in records:
        records[key] = {'attempts': 0, 'errors': 0}
    records[key]['attempts'] += 1
    if not is_correct:
        records[key]['errors'] += 1
    save_records(records)


# ============ 加权随机抽题 ============

def weighted_random_pick(questions, records):
    """根据错误次数和错误率进行加权随机抽题"""
    weights = []
    for q in questions:
        rec = get_record(records, q['id'])
        attempts = rec['attempts']
        errors = rec['errors']
        error_rate = errors / attempts if attempts > 0 else 0

        # 基础权重
        weight = 1.0

        # 错误次数越多权重越大
        if errors >= 5:
            weight += 3.0
        elif errors >= 3:
            weight += 2.0
        elif errors >= 1:
            weight += 1.0

        # 错误率高的加权
        if attempts > 0:
            if error_rate > 0.5:
                weight += 3.0
            elif error_rate > 0.3:
                weight += 1.5
            elif error_rate > 0:
                weight += 0.5

        # 从未做过的题目也适当提高权重
        if attempts == 0:
            weight += 0.5

        weights.append(weight)

    # 加权随机选择
    total = sum(weights)
    r = random.uniform(0, total)
    cumulative = 0
    for i, w in enumerate(weights):
        cumulative += w
        if r <= cumulative:
            return questions[i]
    return questions[-1]


# ============ GUI 模块 ============

class QuizApp:
    def __init__(self, root, questions, source_path=''):
        self.root = root
        self.questions = questions
        self.question_map = {q['id']: q for q in questions}
        # 兼容单文件字符串与多文件显示标签。
        self.source_path = source_path
        if isinstance(source_path, (list, tuple)):
            source_name = f'多文件({len(source_path)})'
        else:
            source_name = os.path.basename(source_path) if source_path else '未命名'
        self.source_name = source_name
        self.records = load_records()
        self.current_q = None
        self.selected = set()
        self.submitted = False
        self.answer_revealed = False
        self.option_buttons = {}
        self.judge_buttons = []
        self.keyboard_var = tk.StringVar()
        self.recent_signatures = []
        self.recent_signature_limit = 6
        self.duplicate_groups = self._build_duplicate_groups()
        self.duplicate_signature_set = self._build_duplicate_signature_set()

        self.root.title("刷题工具")
        self.ui_scale = self._get_ui_scale()
        self._apply_adaptive_window_geometry()
        self.root.configure(bg='#f5f5f5')
        min_w = max(760, int(self.window_width * 0.72))
        min_h = max(560, int(self.window_height * 0.72))
        self.root.minsize(min_w, min_h)

        # 字体
        self.title_font = tkfont.Font(family='Microsoft YaHei', size=self._scale_font(16), weight='bold')
        self.text_font = tkfont.Font(family='Microsoft YaHei', size=self._scale_font(13))
        self.option_font = tkfont.Font(family='Microsoft YaHei', size=self._scale_font(12))
        self.small_font = tkfont.Font(family='Microsoft YaHei', size=self._scale_font(10))
        self.btn_font = tkfont.Font(family='Microsoft YaHei', size=self._scale_font(11), weight='bold')

        self.question_wrap = max(680, int(self.window_width * 0.84))

        self.build_ui()
        self.root.bind('<Return>', self._on_global_enter)
        self.root.after(120, self._refresh_layout)
        self.show_welcome()

    def _get_ui_scale(self):
        screen_w = self.root.winfo_screenwidth()
        screen_h = self.root.winfo_screenheight()
        scale_w = screen_w / 1920
        scale_h = screen_h / 1080
        return max(1.0, min(1.5, (scale_w + scale_h) / 2))

    def _scale_font(self, base_size):
        return max(10, int(round(base_size * self.ui_scale)))

    def _apply_adaptive_window_geometry(self):
        screen_w = self.root.winfo_screenwidth()
        screen_h = self.root.winfo_screenheight()

        self.window_width = max(920, min(1600, int(screen_w * 0.78)))
        self.window_height = max(720, min(1100, int(screen_h * 0.84)))

        x = max(0, (screen_w - self.window_width) // 2)
        y = max(0, (screen_h - self.window_height) // 2)
        self.root.geometry(f"{self.window_width}x{self.window_height}+{x}+{y}")

    def build_ui(self):
        # 顶部信息栏
        top_frame = tk.Frame(self.root, bg='#2c3e50', height=50)
        top_frame.pack(fill='x')
        top_frame.pack_propagate(False)

        self.info_label = tk.Label(
            top_frame,
            text=f"题库：{self.source_name} | 共 {len(self.questions)} 题",
            font=self.small_font, fg='white', bg='#2c3e50'
        )
        self.info_label.pack(side='left', padx=15, pady=10)

        self.stats_label = tk.Label(
            top_frame, text="",
            font=self.small_font, fg='#ecf0f1', bg='#2c3e50'
        )
        self.stats_label.pack(side='right', padx=15, pady=10)

        # 主内容区域（使用 Canvas + Scrollbar 实现滚动）
        main_container = tk.Frame(self.root, bg='#f5f5f5')
        main_container.pack(fill='both', expand=True, padx=20, pady=10)

        self.canvas = tk.Canvas(main_container, bg='#f5f5f5', highlightthickness=0)
        self.scrollbar = tk.Scrollbar(main_container, orient='vertical', command=self.canvas.yview)
        self.scrollable_frame = tk.Frame(self.canvas, bg='#f5f5f5')

        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )
        self.canvas_window = self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)

        self.canvas.pack(side='left', fill='both', expand=True)
        self.scrollbar.pack(side='right', fill='y')

        # 让 scrollable_frame 宽度跟随 canvas
        self.canvas.bind('<Configure>', self._on_canvas_configure)

        # 鼠标滚轮绑定
        self.canvas.bind_all("<MouseWheel>", self._on_mousewheel)
        self.canvas.bind_all("<Button-4>", self._on_mousewheel_linux)
        self.canvas.bind_all("<Button-5>", self._on_mousewheel_linux)

        # 题号标签
        self.q_number_label = tk.Label(
            self.scrollable_frame, text="", font=self.title_font,
            fg='#2c3e50', bg='#f5f5f5', anchor='w'
        )
        self.q_number_label.pack(fill='x', pady=(10, 2))

        # 题目类型标签
        self.q_type_label = tk.Label(
            self.scrollable_frame, text="", font=self.small_font,
            fg='#7f8c8d', bg='#f5f5f5', anchor='w'
        )
        self.q_type_label.pack(fill='x', pady=(0, 5))

        # 题目文本
        self.q_text_label = tk.Label(
            self.scrollable_frame, text="", font=self.text_font,
            fg='#2c3e50', bg='#ffffff', anchor='w', justify='left',
            wraplength=self.question_wrap, padx=15, pady=15, relief='ridge', bd=1
        )
        self.q_text_label.pack(fill='x', pady=(0, 15))

        # 选项容器
        self.options_frame = tk.Frame(self.scrollable_frame, bg='#f5f5f5')
        self.options_frame.pack(fill='x', pady=5)

        # 结果显示
        self.result_label = tk.Label(
            self.scrollable_frame, text="", font=self.text_font,
            fg='#27ae60', bg='#f5f5f5', anchor='w', justify='left',
            wraplength=self.question_wrap
        )
        self.result_label.pack(fill='x', pady=10)

        # 历史记录显示
        self.history_label = tk.Label(
            self.scrollable_frame, text="", font=self.small_font,
            fg='#95a5a6', bg='#f5f5f5', anchor='w'
        )
        self.history_label.pack(fill='x', pady=(0, 5))

        # 底部按钮栏
        btn_frame = tk.Frame(self.root, bg='#ecf0f1', height=60)
        btn_frame.pack(fill='x', side='bottom')
        btn_frame.pack_propagate(False)

        self.submit_btn = tk.Button(
            btn_frame, text="提交答案", font=self.btn_font,
            bg='#3498db', fg='white', activebackground='#2980b9',
            relief='flat', padx=20, pady=8, command=self.submit_answer
        )
        self.submit_btn.pack(side='left', padx=20, pady=10)

        self.next_btn = tk.Button(
            btn_frame, text="下一题 ▶", font=self.btn_font,
            bg='#2ecc71', fg='white', activebackground='#27ae60',
            relief='flat', padx=20, pady=8, command=self.next_question
        )
        self.next_btn.pack(side='left', padx=5, pady=10)

        # 键盘输入区：支持 ABC 选项输入与 t/f 主观自评。
        input_frame = tk.Frame(btn_frame, bg='#ecf0f1')
        input_frame.pack(side='left', padx=(12, 0), pady=10)

        self.keyboard_hint_label = tk.Label(
            input_frame,
            text='键盘输入：A/B/C 或 ABC；主观题用 t/f',
            font=self.small_font,
            fg='#34495e',
            bg='#ecf0f1'
        )
        self.keyboard_hint_label.pack(side='left', padx=(0, 8))

        self.keyboard_entry = tk.Entry(
            input_frame,
            textvariable=self.keyboard_var,
            font=self.option_font,
            width=16,
            relief='groove'
        )
        self.keyboard_entry.pack(side='left', padx=(0, 6))
        self.keyboard_entry.bind('<Return>', self._on_entry_enter)

        self.reset_btn = tk.Button(
            btn_frame, text="重置记录", font=self.small_font,
            bg='#e74c3c', fg='white', activebackground='#c0392b',
            relief='flat', padx=10, pady=5, command=self.reset_records
        )
        self.reset_btn.pack(side='right', padx=20, pady=10)

        self.freq_btn = tk.Button(
            btn_frame, text="考频统计", font=self.small_font,
            bg='#8e44ad', fg='white', activebackground='#7d3c98',
            relief='flat', padx=10, pady=5, command=self.show_frequency_stats
        )
        self.freq_btn.pack(side='right', padx=(0, 8), pady=10)

    def _on_canvas_configure(self, event):
        self.canvas.itemconfig(self.canvas_window, width=event.width)
        self._refresh_layout()

    def _refresh_layout(self):
        canvas_w = self.canvas.winfo_width()
        if canvas_w <= 1:
            canvas_w = max(600, self.root.winfo_width() - 90)

        content_width = max(420, canvas_w - 45)
        self.q_text_label.config(wraplength=content_width)
        self.result_label.config(wraplength=content_width)

        option_wrap = max(380, content_width - 36)
        for btn in self.option_buttons.values():
            btn.config(wraplength=option_wrap)

    def _on_mousewheel(self, event):
        self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def _on_mousewheel_linux(self, event):
        if event.num == 4:
            self.canvas.yview_scroll(-1, "units")
        elif event.num == 5:
            self.canvas.yview_scroll(1, "units")

    def show_welcome(self):
        self.q_number_label.config(text="欢迎使用刷题工具！")
        self.q_type_label.config(text="")
        duplicate_hint = ''
        if self.duplicate_groups:
            dup_questions = sum(len(g) for g in self.duplicate_groups)
            duplicate_hint = f"\n\n检测到重复题组 {len(self.duplicate_groups)} 组（共 {dup_questions} 题），系统会自动尽量避免短时间重复抽到同组题。"

        self.q_text_label.config(
            text=f"题库已加载 {len(self.questions)} 道题目。\n\n"
                 f"点击「下一题」开始刷题！\n\n"
                 f"系统会根据你的错误率自动加权抽题，\n"
                 f"错得越多的题越容易被抽到哦～"
                 f"{duplicate_hint}"
        )
        self.result_label.config(text="")
        self.history_label.config(text="")
        self.keyboard_var.set('')
        self.keyboard_hint_label.config(text='键盘输入：A/B/C 或 ABC；主观题用 t/f')
        self.submit_btn.config(state='disabled')
        self.update_stats()

    def update_stats(self):
        total_attempts = sum(r['attempts'] for r in self.records.values())
        total_errors = sum(r['errors'] for r in self.records.values())
        attempted = len([r for r in self.records.values() if r['attempts'] > 0])
        acc = (1 - total_errors / total_attempts) * 100 if total_attempts > 0 else 0
        self.stats_label.config(
            text=f"已做: {attempted}/{len(self.questions)} | "
                 f"总答题: {total_attempts} | 总正确率: {acc:.1f}%"
        )

    def next_question(self):
        self.submitted = False
        self.answer_revealed = False
        self.selected = set()
        self.keyboard_var.set('')
        picked = weighted_random_pick(self.questions, self.records)

        # 若命中近期重复题组，尝试重抽，减少“同题反复出现”的体感。
        if len(self.questions) > 1 and self.duplicate_groups:
            for _ in range(30):
                if not self._is_recent_duplicate_pick(picked):
                    break
                candidate = weighted_random_pick(self.questions, self.records)
                if not self._is_recent_duplicate_pick(candidate):
                    picked = candidate
                    break
                picked = candidate

        self.current_q = picked
        sig = self._question_signature(self.current_q)
        self.recent_signatures.append(sig)
        if len(self.recent_signatures) > self.recent_signature_limit:
            self.recent_signatures = self.recent_signatures[-self.recent_signature_limit:]
        self.display_question()

    def display_question(self):
        q = self.current_q
        type_map = {
            'single': '【单选题】',
            'multi': '【多选题】',
            'judge': '【判断题】',
            'blank': '【填空题】',
            'short': '【简答题】'
        }

        self.q_number_label.config(text=f"第 {q['id']} 题")
        self.q_type_label.config(text=type_map.get(q['type'], ''))
        display_text = q['text']
        if q.get('type') == 'blank':
            display_text = _mask_blank_question_text(q['text'], q.get('answer', ''))
        self.q_text_label.config(text=display_text)
        self.result_label.config(text="")

        # 显示历史记录
        rec = get_record(self.records, q['id'])
        if rec['attempts'] > 0:
            rate = rec['errors'] / rec['attempts'] * 100
            self.history_label.config(
                text=f"历史记录：答过 {rec['attempts']} 次，错误 {rec['errors']} 次，错误率 {rate:.0f}%"
            )
        else:
            self.history_label.config(text="历史记录：首次作答")

        # 清除旧选项
        for widget in self.options_frame.winfo_children():
            widget.destroy()
        self.option_buttons = {}
        self.judge_buttons = []

        if q['type'] in ('single', 'multi', 'judge'):
            # 客观题：创建选项按钮
            sorted_keys = sorted(q['options'].keys())
            for key in sorted_keys:
                btn = tk.Button(
                    self.options_frame,
                    text=f"  {key}. {q['options'][key]}",
                    font=self.option_font,
                    bg='white', fg='#2c3e50',
                    activebackground='#d5e8d4',
                    relief='ridge', bd=1,
                    anchor='w', justify='left',
                    padx=15, pady=8,
                    wraplength=700,
                    command=lambda k=key: self.toggle_option(k)
                )
                btn.pack(fill='x', pady=3, ipady=3)
                self.option_buttons[key] = btn

            self.submit_btn.config(text='提交答案', state='normal')
            if q['type'] == 'multi':
                self.keyboard_hint_label.config(text='键盘输入：如 ABC；回车提交，已判分后回车下一题')
            else:
                self.keyboard_hint_label.config(text='键盘输入：如 A 或 B；回车提交，已判分后回车下一题')
        else:
            # 主观题：先不显示任何作答选项，先看答案再自判
            self.result_label.config(
                text='请先自己作答，点击「显示正确答案」后再进行自评。',
                fg='#7f8c8d'
            )
            self.submit_btn.config(text='显示正确答案', state='normal')
            self.keyboard_hint_label.config(text='主观题：先回车显示答案，再输入 t/f 回车自评；再回车下一题')

        self._refresh_layout()
        self.keyboard_entry.focus_set()

        # 滚动到顶部
        self.canvas.yview_moveto(0)

    def toggle_option(self, key):
        if self.submitted:
            return

        q = self.current_q
        if q['type'] == 'single' or q['type'] == 'judge':
            # 单选：清除其他选中
            self.selected = {key}
            for k, btn in self.option_buttons.items():
                if k == key:
                    btn.config(bg='#3498db', fg='white')
                else:
                    btn.config(bg='white', fg='#2c3e50')
        else:
            # 多选：切换选中状态
            if key in self.selected:
                self.selected.discard(key)
                self.option_buttons[key].config(bg='white', fg='#2c3e50')
            else:
                self.selected.add(key)
                self.option_buttons[key].config(bg='#3498db', fg='white')

    def submit_answer(self):
        if self.submitted:
            return

        q = self.current_q

        if q['type'] in ('blank', 'short'):
            if self.answer_revealed:
                return

            self.answer_revealed = True
            self.result_label.config(
                text=f"参考答案：{q['answer']}\n\n请根据你的作答进行自评：",
                fg='#2c3e50'
            )

            correct_btn = tk.Button(
                self.options_frame,
                text='我答对了',
                font=self.btn_font,
                bg='#2ecc71', fg='white',
                activebackground='#27ae60',
                relief='flat', padx=15, pady=8,
                command=lambda: self.submit_subjective_result(True)
            )
            correct_btn.pack(side='left', padx=(0, 10), pady=5)

            wrong_btn = tk.Button(
                self.options_frame,
                text='我答错了',
                font=self.btn_font,
                bg='#e74c3c', fg='white',
                activebackground='#c0392b',
                relief='flat', padx=15, pady=8,
                command=lambda: self.submit_subjective_result(False)
            )
            wrong_btn.pack(side='left', pady=5)

            self.judge_buttons = [correct_btn, wrong_btn]
            self.submit_btn.config(state='disabled')
            self.keyboard_hint_label.config(text='输入 t(对)/f(错) 后回车自评，已记录后回车下一题')
            return

        if not self.selected:
            messagebox.showwarning("提示", "请先选择一个答案！")
            return

        if q['type'] in ('single', 'multi', 'judge') and not q.get('answer'):
            messagebox.showwarning("提示", "本题未识别出标准答案，暂无法自动判分。")
            return

        self.submitted = True
        correct = set(q['answer'])
        is_correct = self.selected == correct

        # 更新记录
        update_record(self.records, q['id'], is_correct)
        self.update_stats()

        # 高亮显示正确/错误选项
        for k, btn in self.option_buttons.items():
            if k in correct:
                btn.config(bg='#27ae60', fg='white')  # 正确选项绿色
            elif k in self.selected and k not in correct:
                btn.config(bg='#e74c3c', fg='white')  # 错选红色
            else:
                btn.config(bg='#ecf0f1', fg='#95a5a6')  # 未选灰色

        if is_correct:
            self.result_label.config(
                text="✓ 回答正确！", fg='#27ae60'
            )
        else:
            self.result_label.config(
                text=f"✗ 回答错误！正确答案是：{''.join(sorted(correct))}",
                fg='#e74c3c'
            )

        # 更新历史显示
        rec = get_record(self.records, q['id'])
        rate = rec['errors'] / rec['attempts'] * 100
        self.history_label.config(
            text=f"历史记录：答过 {rec['attempts']} 次，错误 {rec['errors']} 次，错误率 {rate:.0f}%"
        )

        self.submit_btn.config(state='disabled')
        self.keyboard_hint_label.config(text='已判分，按回车进入下一题')

    def submit_subjective_result(self, is_correct):
        if self.submitted:
            return

        self.submitted = True
        q = self.current_q

        update_record(self.records, q['id'], is_correct)
        self.update_stats()

        for btn in self.judge_buttons:
            btn.config(state='disabled')

        if is_correct:
            self.result_label.config(
                text=f"参考答案：{q['answer']}\n\n✓ 已记录：你判定本题答对。",
                fg='#27ae60'
            )
        else:
            self.result_label.config(
                text=f"参考答案：{q['answer']}\n\n✗ 已记录：你判定本题答错。",
                fg='#e74c3c'
            )

        rec = get_record(self.records, q['id'])
        rate = rec['errors'] / rec['attempts'] * 100
        self.history_label.config(
            text=f"历史记录：答过 {rec['attempts']} 次，错误 {rec['errors']} 次，错误率 {rate:.0f}%"
        )

        self.submit_btn.config(state='disabled')
        self.keyboard_hint_label.config(text='已记录，按回车进入下一题')

    def reset_records(self):
        if messagebox.askyesno("确认", "确定要重置所有错误记录吗？\n此操作不可撤销！"):
            self.records = {}
            save_records(self.records)
            self.update_stats()
            messagebox.showinfo("完成", "所有记录已重置。")

    def show_frequency_stats(self):
        """展示题目考频统计（基于作答记录）。"""
        win = tk.Toplevel(self.root)
        win.title('考频统计')
        win_w = max(900, int(self.root.winfo_width() * 0.9))
        win_h = max(560, int(self.root.winfo_height() * 0.8))
        x = self.root.winfo_x() + max(0, (self.root.winfo_width() - win_w) // 2)
        y = self.root.winfo_y() + max(0, (self.root.winfo_height() - win_h) // 2)
        win.geometry(f'{win_w}x{win_h}+{x}+{y}')
        win.configure(bg='#f5f5f5')
        win.transient(self.root)

        # 汇总数据
        rows = []
        attempted_count = 0
        total_attempts = 0
        total_errors = 0
        type_map = {
            'single': '单选',
            'multi': '多选',
            'judge': '判断',
            'blank': '填空',
            'short': '简答'
        }

        for q in self.questions:
            rec = get_record(self.records, q['id'])
            attempts = rec['attempts']
            errors = rec['errors']
            error_rate = (errors / attempts * 100) if attempts > 0 else 0.0
            if attempts > 0:
                attempted_count += 1
            total_attempts += attempts
            total_errors += errors

            rows.append({
                'id': q['id'],
                'type': type_map.get(q.get('type', ''), q.get('type', '')),
                'attempts': attempts,
                'errors': errors,
                'error_rate': error_rate,
                'text': str(q.get('text', '')).replace('\n', ' ').strip()
            })

        total_questions = len(self.questions)
        overall_error_rate = (total_errors / total_attempts * 100) if total_attempts > 0 else 0.0

        top_frame = tk.Frame(win, bg='#2c3e50', height=56)
        top_frame.pack(fill='x')
        top_frame.pack_propagate(False)

        summary_label = tk.Label(
            top_frame,
            text=(
                f'总题数 {total_questions}  |  已作答 {attempted_count}  |  总作答次数 {total_attempts}  '
                f'|  总错误次数 {total_errors}  |  总错误率 {overall_error_rate:.1f}%'
            ),
            font=self.small_font,
            fg='white',
            bg='#2c3e50',
            anchor='w'
        )
        summary_label.pack(fill='x', padx=12, pady=16)

        if self.duplicate_groups:
            dup_questions = sum(len(g) for g in self.duplicate_groups)
            dup_label = tk.Label(
                win,
                text=f'重复题检测：{len(self.duplicate_groups)} 组（共 {dup_questions} 题），系统抽题时会自动规避短时间重复同组。',
                font=self.small_font,
                fg='#8e44ad',
                bg='#f5f5f5',
                anchor='w'
            )
            dup_label.pack(fill='x', padx=12, pady=(6, 0))

        control_frame = tk.Frame(win, bg='#f5f5f5')
        control_frame.pack(fill='x', padx=12, pady=(10, 6))

        tk.Label(
            control_frame, text='排序方式：',
            font=self.small_font, bg='#f5f5f5', fg='#2c3e50'
        ).pack(side='left')

        sort_var = tk.StringVar(value='按作答次数')
        sort_combo = ttk.Combobox(
            control_frame,
            state='readonly',
            textvariable=sort_var,
            values=['按作答次数', '按错误次数', '按错误率', '按题号'],
            width=14
        )
        sort_combo.pack(side='left', padx=(6, 10))

        only_attempted_var = tk.BooleanVar(value=True)
        only_attempted_cb = tk.Checkbutton(
            control_frame,
            text='仅看已作答题目',
            variable=only_attempted_var,
            bg='#f5f5f5',
            fg='#2c3e50',
            font=self.small_font,
            activebackground='#f5f5f5'
        )
        only_attempted_cb.pack(side='left')

        body = tk.Frame(win, bg='#f5f5f5')
        body.pack(fill='both', expand=True, padx=12, pady=(4, 10))

        columns = ('rank', 'id', 'type', 'attempts', 'errors', 'error_rate', 'text')
        tree = ttk.Treeview(body, columns=columns, show='headings')
        tree.heading('rank', text='排名')
        tree.heading('id', text='题号')
        tree.heading('type', text='题型')
        tree.heading('attempts', text='作答次数')
        tree.heading('errors', text='错误次数')
        tree.heading('error_rate', text='错误率')
        tree.heading('text', text='题干预览')

        tree.column('rank', width=56, anchor='center')
        tree.column('id', width=68, anchor='center')
        tree.column('type', width=72, anchor='center')
        tree.column('attempts', width=88, anchor='center')
        tree.column('errors', width=88, anchor='center')
        tree.column('error_rate', width=84, anchor='center')
        tree.column('text', width=620, anchor='w')

        yscroll = ttk.Scrollbar(body, orient='vertical', command=tree.yview)
        xscroll = ttk.Scrollbar(body, orient='horizontal', command=tree.xview)
        tree.configure(yscrollcommand=yscroll.set, xscrollcommand=xscroll.set)

        tree.grid(row=0, column=0, sticky='nsew')
        yscroll.grid(row=0, column=1, sticky='ns')
        xscroll.grid(row=1, column=0, sticky='ew')
        body.grid_columnconfigure(0, weight=1)
        body.grid_rowconfigure(0, weight=1)

        def apply_sort(items):
            mode = sort_var.get()
            if mode == '按作答次数':
                return sorted(items, key=lambda x: (-x['attempts'], -x['errors'], x['id']))
            if mode == '按错误次数':
                return sorted(items, key=lambda x: (-x['errors'], -x['attempts'], x['id']))
            if mode == '按错误率':
                return sorted(items, key=lambda x: (-x['error_rate'], -x['attempts'], x['id']))
            return sorted(items, key=lambda x: x['id'])

        def refresh_table(_event=None):
            for item in tree.get_children():
                tree.delete(item)

            display_rows = rows
            if only_attempted_var.get():
                display_rows = [r for r in display_rows if r['attempts'] > 0]

            display_rows = apply_sort(display_rows)
            for i, r in enumerate(display_rows, 1):
                text_preview = r['text']
                if len(text_preview) > 90:
                    text_preview = text_preview[:90] + '...'
                tree.insert('', 'end', values=(
                    i,
                    r['id'],
                    r['type'],
                    r['attempts'],
                    r['errors'],
                    f"{r['error_rate']:.0f}%",
                    text_preview
                ))

        sort_combo.bind('<<ComboboxSelected>>', refresh_table)
        only_attempted_cb.config(command=refresh_table)
        refresh_table()

    def _normalize_keyboard_text(self, text):
        normalized = (text or '').strip().upper()
        return normalized.translate(str.maketrans('ＡＢＣＤＥＦＧＨ，。、；：　', 'ABCDEFGH,,,,  '))

    def _question_signature(self, q):
        """生成题目归一签名，用于重复题检测。"""
        text = str(q.get('text', '') or '')
        text = re.sub(r'（\s*\d+\s*）\s*[_＿﹍]+', '（）', text)
        text = re.sub(r'[_＿﹍]+', '', text)
        text = re.sub(r'[\s，,。；;：:、（）()\[\]【】]+', '', text)

        options = q.get('options') or {}
        option_sig = []
        for k in sorted(options.keys()):
            v = re.sub(r'\s+', '', str(options.get(k, '') or ''))
            option_sig.append(f'{k}:{v}')

        q_type = q.get('type', '')
        return (q_type, text, '|'.join(option_sig))

    def _build_duplicate_groups(self):
        sig_map = {}
        for q in self.questions:
            sig = self._question_signature(q)
            sig_map.setdefault(sig, []).append(q['id'])

        groups = [ids for ids in sig_map.values() if len(ids) >= 2]
        groups.sort(key=lambda x: (-len(x), x[0]))
        return groups

    def _is_recent_duplicate_pick(self, q):
        sig = self._question_signature(q)
        # 仅当该签名属于重复题组，且近期出现过，才判定为重复抽到。
        return sig in self.duplicate_signature_set and sig in self.recent_signatures

    def _build_duplicate_signature_set(self):
        sigs = set()
        for g in self.duplicate_groups:
            for qid in g:
                q = self.question_map.get(qid)
                if q:
                    sigs.add(self._question_signature(q))
        return sigs

    def _select_objective_by_keyboard(self, token):
        if not self.current_q or self.submitted:
            return False

        q = self.current_q
        if q['type'] not in ('single', 'multi', 'judge'):
            return False

        valid_keys = sorted(q['options'].keys())
        letters = [ch for ch in token if ch in valid_keys]
        if not letters:
            return False

        if q['type'] in ('single', 'judge'):
            self.selected = {letters[-1]}
        else:
            self.selected = set(letters)

        for k, btn in self.option_buttons.items():
            if k in self.selected:
                btn.config(bg='#3498db', fg='white')
            else:
                btn.config(bg='white', fg='#2c3e50')
        return True

    def _submit_subjective_by_keyboard(self, token):
        if not self.current_q or self.submitted:
            return False

        q = self.current_q
        if q['type'] not in ('blank', 'short'):
            return False

        if not self.answer_revealed:
            return False

        true_tokens = {'T', 'TRUE', 'Y', 'YES', '对', '正确'}
        false_tokens = {'F', 'FALSE', 'N', 'NO', '错', '错误'}
        if token in true_tokens:
            self.submit_subjective_result(True)
            return True
        if token in false_tokens:
            self.submit_subjective_result(False)
            return True
        return False

    def _process_keyboard_enter(self):
        if self.current_q is None:
            self.next_question()
            return

        token = self._normalize_keyboard_text(self.keyboard_var.get())
        self.keyboard_var.set('')

        if token:
            q_type = self.current_q['type']
            if q_type in ('single', 'multi', 'judge'):
                if not self._select_objective_by_keyboard(token):
                    self.result_label.config(text='未识别到有效选项，请输入题目存在的字母（如 A/ABC）。', fg='#e67e22')
                    return
                if not self.submitted:
                    self.submit_answer()
                return

            if q_type in ('blank', 'short'):
                if self._submit_subjective_by_keyboard(token):
                    return
                self.result_label.config(text='主观题请输入 t/f 后回车（t=答对，f=答错）。', fg='#e67e22')
                return

        # 无输入时：未提交则提交/显示答案；已提交则下一题。
        if not self.submitted:
            self.submit_answer()
        else:
            self.next_question()

    def _on_entry_enter(self, _event=None):
        self._process_keyboard_enter()
        return 'break'

    def _on_global_enter(self, event=None):
        widget = self.root.focus_get()
        if widget is self.keyboard_entry:
            return
        self._process_keyboard_enter()


# ============ 主程序 ============

def main():
    _enable_windows_high_dpi()

    # 确定脚本目录
    script_dir = os.path.dirname(os.path.abspath(__file__))

    # 文件选择阶段使用独立隐藏窗口，避免主窗口状态异常导致不显示
    selector = tk.Tk()
    selector.withdraw()
    selector.update_idletasks()

    # 手动选择题库文件（支持多选）
    selected_paths = filedialog.askopenfilenames(
        title='请选择一个或多个题库文件',
        initialdir=script_dir,
        filetypes=[
            ('题库文件', '*.txt *.pdf *.doc *.docx'),
            ('文本文件', '*.txt'),
            ('PDF 文件', '*.pdf'),
            ('Word 文件', '*.doc *.docx'),
            ('所有文件', '*.*')
        ],
        parent=selector
    )
    selector.destroy()

    file_paths = list(selected_paths)
    if not file_paths:
        messagebox.showinfo('已取消', '未选择题库文件，程序将退出。')
        return

    print(f"正在解析题库，共 {len(file_paths)} 个文件。")

    if len(file_paths) == 1:
        txt_path = file_paths[0]
        try:
            candidates = build_parse_candidates(txt_path)
        except Exception as e:
            messagebox.showerror('解析失败', f'文件解析失败：\n{e}')
            return

        if not candidates or not candidates[0][1]:
            messagebox.showerror('错误', '未能解析出任何题目，请检查题库格式。')
            return

        source_label = txt_path
    else:
        merged_questions = []
        failed_files = []

        for path in file_paths:
            try:
                file_candidates = build_parse_candidates(path)
                if not file_candidates or not file_candidates[0][1]:
                    failed_files.append((path, '未解析出题目'))
                    continue

                file_questions = file_candidates[0][1]
                for q in file_questions:
                    copied = dict(q)
                    copied['source_file'] = os.path.basename(path)
                    merged_questions.append(copied)
                print(f"已解析：{os.path.basename(path)} -> {len(file_questions)} 题")
            except Exception as e:
                failed_files.append((path, str(e)))

        if not merged_questions:
            detail = '\n'.join(f"- {os.path.basename(p)}: {msg}" for p, msg in failed_files) if failed_files else '未知错误'
            messagebox.showerror('解析失败', f'所有文件均解析失败：\n{detail}')
            return

        for idx, q in enumerate(merged_questions, 1):
            q['id'] = idx

        if failed_files:
            detail = '\n'.join(f"- {os.path.basename(p)}: {msg}" for p, msg in failed_files[:8])
            messagebox.showwarning('部分文件解析失败', f'以下文件未成功导入（其余文件已合并）：\n{detail}')

        candidates = [
            (
                f'自动识别（多文件合并，{len(file_paths)}个）',
                merged_questions,
                '每个文件采用自动识别方案后合并为同一题库'
            )
        ]
        source_label = [os.path.basename(p) for p in file_paths]

    # 主窗口单独创建，确保可见
    root = tk.Tk()
    root.title('刷题工具')
    root.lift()
    root.focus_force()

    # 导入预览窗口：先确认解析结果再进入刷题
    preview_source = file_paths[0] if len(file_paths) == 1 else f'多文件导入（{len(file_paths)}个）'
    selected_questions = show_import_preview(root, candidates, preview_source)
    if not selected_questions:
        root.destroy()
        return

    # 切换工作目录到脚本目录（使 error_record.json 保存在同一位置）
    os.chdir(script_dir)

    app = QuizApp(root, selected_questions, source_path=source_label)
    root.mainloop()


def show_import_preview(root, candidates, source_path):
    """展示解析预览，用户确认后开始刷题。"""
    win = tk.Toplevel(root)
    win.title('题库导入预览')
    screen_w = root.winfo_screenwidth()
    screen_h = root.winfo_screenheight()
    win_w = max(900, min(1700, int(screen_w * 0.82)))
    win_h = max(580, min(1100, int(screen_h * 0.8)))
    x = max(0, (screen_w - win_w) // 2)
    y = max(0, (screen_h - win_h) // 2)
    win.geometry(f'{win_w}x{win_h}+{x}+{y}')
    win.minsize(max(850, int(win_w * 0.75)), max(520, int(win_h * 0.72)))
    win.configure(bg='#f5f5f5')
    win.transient(root)
    win.grab_set()

    top = tk.Frame(win, bg='#2c3e50', height=52)
    top.pack(fill='x')
    top.pack_propagate(False)

    top_label = tk.Label(
        top, text='', fg='white', bg='#2c3e50',
        font=('Microsoft YaHei', 10), anchor='w'
    )
    top_label.pack(fill='x', padx=12, pady=14)

    body = tk.Frame(win, bg='#f5f5f5')
    body.pack(fill='both', expand=True, padx=12, pady=10)

    columns = ('id', 'type', 'answer', 'text')
    tree = ttk.Treeview(body, columns=columns, show='headings', height=22)
    tree.heading('id', text='题号')
    tree.heading('type', text='题型')
    tree.heading('answer', text='答案')
    tree.heading('text', text='题干预览')
    tree.column('id', width=70, anchor='center')
    tree.column('type', width=90, anchor='center')
    tree.column('answer', width=170, anchor='w')
    tree.column('text', width=600, anchor='w')

    type_name = {
        'single': '单选',
        'multi': '多选',
        'judge': '判断',
        'blank': '填空',
        'short': '简答'
    }

    selected_idx = tk.IntVar(value=0)

    strategy_frame = tk.Frame(win, bg='#f5f5f5')
    strategy_frame.pack(fill='x', padx=12, pady=(0, 8))

    tk.Label(
        strategy_frame,
        text='识别方案：',
        bg='#f5f5f5', fg='#2c3e50', font=('Microsoft YaHei', 10)
    ).pack(side='left')

    strategy_names = [name for name, _, _ in candidates]
    strategy_combo = ttk.Combobox(
        strategy_frame,
        values=strategy_names,
        state='readonly',
        width=34
    )
    strategy_combo.current(0)
    strategy_combo.pack(side='left', padx=(6, 10))

    strategy_desc_label = tk.Label(
        strategy_frame,
        text='',
        bg='#f5f5f5', fg='#7f8c8d', font=('Microsoft YaHei', 9), anchor='w'
    )
    strategy_desc_label.pack(side='left', fill='x', expand=True)

    def fill_tree(questions):
        for item in tree.get_children():
            tree.delete(item)

        for q in questions:
            if isinstance(q.get('answer'), list):
                answer_preview = ''.join(q['answer'])
            else:
                answer_preview = str(q.get('answer', ''))
            if not answer_preview.strip():
                answer_preview = '（未识别）'
            answer_preview = answer_preview.replace('\n', ' / ').strip()
            text_preview = str(q.get('text', '')).replace('\n', ' ').strip()
            if len(text_preview) > 80:
                text_preview = text_preview[:80] + '...'
            if len(answer_preview) > 32:
                answer_preview = answer_preview[:32] + '...'

            tree.insert('', 'end', values=(
                q.get('id', ''),
                type_name.get(q.get('type'), q.get('type', '')),
                answer_preview,
                text_preview
            ))

    def refresh_summary(idx):
        name, questions, desc = candidates[idx]
        type_count = {'single': 0, 'multi': 0, 'judge': 0, 'blank': 0, 'short': 0}
        for q in questions:
            t = q.get('type')
            if t in type_count:
                type_count[t] += 1

        summary_text = (
            f"文件：{os.path.basename(source_path)}  |  方案：{name}  |  共 {len(questions)} 题"
            f"  |  单选 {type_count['single']}  多选 {type_count['multi']}  判断 {type_count['judge']}"
            f"  填空 {type_count['blank']}  简答 {type_count['short']}"
        )
        top_label.config(text=summary_text)
        strategy_desc_label.config(text=desc)
        fill_tree(questions)

    def on_strategy_change(_event=None):
        idx = strategy_combo.current()
        if idx < 0:
            idx = 0
        selected_idx.set(idx)
        refresh_summary(idx)

    strategy_combo.bind('<<ComboboxSelected>>', on_strategy_change)
    refresh_summary(0)

    yscroll = ttk.Scrollbar(body, orient='vertical', command=tree.yview)
    xscroll = ttk.Scrollbar(body, orient='horizontal', command=tree.xview)
    tree.configure(yscrollcommand=yscroll.set, xscrollcommand=xscroll.set)

    tree.grid(row=0, column=0, sticky='nsew')
    yscroll.grid(row=0, column=1, sticky='ns')
    xscroll.grid(row=1, column=0, sticky='ew')
    body.grid_columnconfigure(0, weight=1)
    body.grid_rowconfigure(0, weight=1)

    hint = tk.Label(
        win,
        text='请检查题型与答案是否合理。确认无误后点击「开始刷题」。',
        bg='#f5f5f5', fg='#7f8c8d', font=('Microsoft YaHei', 10)
    )
    hint.pack(fill='x', padx=12, pady=(0, 8))

    result = {'ok': False, 'questions': None}

    btn_bar = tk.Frame(win, bg='#f5f5f5')
    btn_bar.pack(fill='x', padx=12, pady=(0, 12))

    def on_cancel():
        result['ok'] = False
        win.destroy()

    def on_confirm():
        result['ok'] = True
        idx = selected_idx.get()
        if idx < 0 or idx >= len(candidates):
            idx = 0
        result['questions'] = candidates[idx][1]
        win.destroy()

    tk.Button(
        btn_bar, text='取消',
        font=('Microsoft YaHei', 10),
        bg='#e0e0e0', fg='#2c3e50',
        relief='flat', padx=16, pady=6,
        command=on_cancel
    ).pack(side='right', padx=(8, 0))

    tk.Button(
        btn_bar, text='开始刷题',
        font=('Microsoft YaHei', 10, 'bold'),
        bg='#2ecc71', fg='white',
        activebackground='#27ae60',
        relief='flat', padx=16, pady=6,
        command=on_confirm
    ).pack(side='right')

    win.protocol('WM_DELETE_WINDOW', on_cancel)
    root.wait_window(win)
    if result['ok']:
        return result['questions']
    return None


def _enable_windows_high_dpi():
    """启用 Windows 高 DPI 感知，避免缩放模糊。"""
    if os.name != 'nt':
        return

    try:
        ctypes.windll.shcore.SetProcessDpiAwareness(1)
    except Exception:
        try:
            ctypes.windll.user32.SetProcessDPIAware()
        except Exception:
            pass


if __name__ == '__main__':
    main()
